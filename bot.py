import os
import re
import json
import logging
import time
import asyncio
import tempfile
import base64
import email.mime.text
from collections import defaultdict
from datetime import datetime, timedelta, timezone
from anthropic import Anthropic

# ---- Garage trigger detection (Engine agent 2026-05-09) --------------------
_GARAGE_PATTERN_RAILWAY = re.compile(
    r"(алло\s*,?\s*гараж|hello\s+garage|hey\s+garage)",
    re.IGNORECASE | re.UNICODE,
)

_GARAGE_SYSTEM_CONTEXT_RAILWAY = (
    "[GARAGE MODE]\n"
    "You are the Garage specialist agent. Your domains:\n"
    "- Cars: maintenance, specs, diagnostics, repair tips.\n"
    "- Smart home: Home Assistant at http://100.106.1.80:8123, SwitchBot, Tapo, sensors.\n"
    "- Devices: nucbox-m5pro (100.106.1.80), Daniel desktop (100.103.98.25).\n"
    "- Networking: Tailscale mesh, WSL2, Docker, router config.\n"
    "- Servers: Windows services, NSSM, scheduled tasks, processes.\n"
    "Respond concisely. Note: this bot runs on Railway and cannot directly SSH or run local commands,\n"
    "but can advise on all garage topics and draft commands for Daniel to run.\n\n"
)


def _detect_garage_railway(text):
    words = text.strip().split()
    prefix = " ".join(words[:3])
    for candidate in ([prefix, text.strip()] if len(words) > 3 else [text.strip()]):
        m = _GARAGE_PATTERN_RAILWAY.search(candidate)
        if m:
            payload = _GARAGE_PATTERN_RAILWAY.sub("", text.strip(), count=1).strip().lstrip(", .")
            return True, payload
    return False, text.strip()

# ---------------------------------------------------------------------------
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Application, CommandHandler, MessageHandler, CallbackQueryHandler, filters, ContextTypes
import urllib.request
import urllib.error
import urllib.parse
from apscheduler.schedulers.asyncio import AsyncIOScheduler
from apscheduler.triggers.cron import CronTrigger

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger(__name__)

TELEGRAM_TOKEN = os.environ["TELEGRAM_BOT_TOKEN"]
ANTHROPIC_API_KEY = os.environ["ANTHROPIC_API_KEY"]
HUBSPOT_API_KEY = os.environ.get("HUBSPOT_API_KEY", "")
ALLOWED_USERS = os.environ.get("ALLOWED_USERS", "")
MODEL = os.environ.get("CLAUDE_MODEL", "claude-sonnet-4-20250514")
MAX_HISTORY = int(os.environ.get("MAX_HISTORY", "50"))

GOOGLE_CLIENT_ID = os.environ.get("GOOGLE_CLIENT_ID", "")
GOOGLE_CLIENT_SECRET = os.environ.get("GOOGLE_CLIENT_SECRET", "")
GOOGLE_REFRESH_TOKEN = os.environ.get("GOOGLE_REFRESH_TOKEN", "")

client = Anthropic(api_key=ANTHROPIC_API_KEY)
# Conversation history persists in SQLite on the Railway volume at /data so
# git push / redeploys don't wipe context. See db.py.
import db as _convdb  # noqa: E402
_convdb.init_db()
pending_updates: dict[int, dict] = {}
scheduled_jobs: dict[int, list] = defaultdict(list)
scheduler = AsyncIOScheduler()
_google_token_cache = {"token": "", "expires": 0}

# Per-chat asyncio lock so a second message arriving mid-process gets
# acknowledged immediately and queued instead of vanishing into the PTB queue.
_CHAT_LOCKS: dict[int, asyncio.Lock] = {}


def _chat_lock(chat_id: int) -> asyncio.Lock:
    lock = _CHAT_LOCKS.get(chat_id)
    if lock is None:
        lock = asyncio.Lock()
        _CHAT_LOCKS[chat_id] = lock
    return lock


async def _ack_if_locked(msg, chat_id: int) -> None:
    if _chat_lock(chat_id).locked():
        try:
            await msg.reply_text("⏳ Доделаю предыдущее, потом возьмусь за это.")
        except Exception:
            pass


# Claude can ask the bot to attach a file to its reply by ending its response
# with one or more [ATTACH: /full/path/to/file] markers. On Railway the
# container fs is ephemeral, so /tmp is the canonical save location.
ATTACH_RE = re.compile(r"\[ATTACH:\s*([^\]\n]+?)\s*\]")
ATTACH_MAX_BYTES = 50 * 1024 * 1024


def _extract_attachments(text):
    paths = [m.group(1).strip() for m in ATTACH_RE.finditer(text or "")]
    cleaned = ATTACH_RE.sub("", text or "").strip()
    return cleaned, paths


async def _send_attachments(reply_target, paths):
    import os as _os
    for raw in paths:
        try:
            if not _os.path.isfile(raw):
                try:
                    await reply_target.reply_text(
                        f"⚠️ ATTACH file not found: {raw}"
                    )
                except Exception:
                    pass
                continue
            sz = _os.path.getsize(raw)
            if sz > ATTACH_MAX_BYTES:
                try:
                    await reply_target.reply_text(
                        f"⚠️ ATTACH too large for Telegram: {_os.path.basename(raw)} "
                        f"({sz // 1024 // 1024} MB > 50 MB)"
                    )
                except Exception:
                    pass
                continue
            with open(raw, "rb") as fh:
                await reply_target.reply_document(fh, filename=_os.path.basename(raw))
        except Exception as exc:
            try:
                await reply_target.reply_text(f"⚠️ ATTACH failed for {raw}: {exc}")
            except Exception:
                pass


# Lazy faster-whisper model (CPU int8). Default size "small"; override on
# small Railway instances with WHISPER_MODEL_SIZE=tiny|base|small.
_whisper_model = None
_WHISPER_SIZE = os.environ.get("WHISPER_MODEL_SIZE", "small")


def _get_whisper():
    global _whisper_model
    if _whisper_model is None:
        from faster_whisper import WhisperModel
        logger.info(f"Loading Whisper model ({_WHISPER_SIZE}, int8 CPU)...")
        _whisper_model = WhisperModel(
            _WHISPER_SIZE, device="cpu", compute_type="int8", cpu_threads=4
        )
        logger.info("Whisper model loaded.")
    return _whisper_model


def _transcribe_sync(file_path: str) -> tuple[str, str]:
    model = _get_whisper()
    segments, info = model.transcribe(file_path, beam_size=5, vad_filter=True)
    text = " ".join(seg.text.strip() for seg in segments).strip()
    return text, info.language

# Cache: telegram username (lowercase, no @) -> HubSpot contact id
# Populated after successful contact creation to avoid search index delay
created_contacts: dict[str, str] = {}

HUBSPOT_BASE = "https://api.hubapi.com"
LIFECYCLE_STAGES = {
    "subscriber": "Subscriber", "lead": "Lead", "marketingqualifiedlead": "MQL",
    "salesqualifiedlead": "SQL", "opportunity": "Opportunity", "customer": "Customer",
    "evangelist": "Evangelist", "1578580705": "Not interested", "other": "Other",
}
LEAD_STATUSES = {
    "NEW": "New", "OPEN": "Open", "IN_PROGRESS": "In Progress",
    "OPEN_DEAL": "Open Deal", "UNQUALIFIED": "Unqualified",
    "ATTEMPTED_TO_CONTACT": "Attempted to Contact", "CONNECTED": "Connected",
    "BAD_TIMING": "Bad Timing",
}

SYSTEM_PROMPT = """You are a personal AI assistant for Daniel (Danil) Tonkopiy. You know him deeply and serve as his right hand.

== WHO DANIEL IS ==
- Full name: Данил Тонкопий / Daniel Tonkopi. Age 47. From Kyiv, Ukraine.
- Lives at Los Altos, CA (Springwood Apartment Homes) with wife Лия (Leah Altukhova, married May 2020) and cat Basiko.
- Drives 2017 Mini Cooper F56. Phone: +12135649148. Email: dt@visanow.ai, tonkopiy@gmail.com
- Citizenship: Kazakhstan. US EB1 green card holder.
- LinkedIn: "CEO | 3 Exits | $1.4B+ in Deals | Guinness World Record | Stanford"

== CURRENT VENTURES ==
- **CoreElement.AI** (Oct 2025–present) — CEO. AI for geological exploration. "Intelligent Mineral Discovery."
  Team: Anton Ermakov (Chief Scientist), Alibek Jussisbayev (CTO), Zhambyl Suraganov (Corporate Relations).
  Active deals: Kazatomprom, Tau-Ken Samruk (Chairman: Nariman Absametov). Shunayi project under NDA.
  Fundraising: pre-money $12-28M range. Target investors: BEV, DCVC, Rio Tinto Ventures, BHP Ventures.
  Applied to YC Spring 2026. NVIDIA Inception member. Stanford/Berkeley/Nazarbayev partnerships.
  Cold email infrastructure: 139 accounts, 46 domains, 56K verified leads. Campaign launch mid-April 2026.
- **JustiGuide** (acquired VisaNow.AI Jan 2026) — immigration automation. Still active referral relationship with attorney Tatiana Aristova (215-350-7972, immigrationwise.com). EB1: $5K+$1K RFE+$2K approval. EB2-NIW: $2.5K+$1K+$2K.
- **iMigrant** — YouTube channel, 800+ videos on US immigration. @imigrant1
- **Apartment Referral Project** — 122 active leads, not yet monetized. GTM: Facebook, Reddit, university communities, immigrant groups.

== CAREER ARC ==
KazMunayGas (geophysicist, gas station network 6→500+, sold $500M) → best.ua (acquired) → X-Rift (AR game) → Delfast (e-bikes, Guinness Record, $3.6M crowdfunded, 3 US patents, Forbes #1) → In Charge One (EV chargers) → VisaNow.AI (acquired) → CoreElement.AI

== FAMILY — PRIORITY #1 ==
- **Sons Никита (~17) and Макар (~17)** live with ex-wife Оксана in Ukraine. THIS IS THE MOST IMPORTANT TOPIC.
  - Макар: reflective, studies extra math, goes to gym. Birthday ~March 25.
  - Никита: more guarded 1-on-1, opens up in groups. Had taxi accident Dec 2025.
  - **CRITICAL:** Sons face military draft risk in Ukraine. University enrollment = protection. "Границы закрывают, интернет отрубают." Nikita considering permanent move to US.
  - Summer 2026 visit planned (June-July) — Makar declined, Nikita deciding.
  - Daniel pursuing I-130 (family reunification).
- **Sister Анюта** — lives in Lisbon. Kids: Рафаэль (2009), Лея (2012). Family emotional anchor.
- **Лия (Leah)** — wife. 43K Telegram msgs. Immigration crisis nearly destroyed the relationship. Both went to therapy.

== FINANCIAL ==
- Credit scores: Equifax 708 / TransUnion 703
- Total debt: ~$15.3K (auto loan $4.2K + credit cards ~$11.1K)
- 14 active credit cards, 100% on-time payments
- Whole Foods: $500-800/month. Known ice cream problem (4+ pints/week vs sugar elimination goal).
- Aspiration: $20M net worth. Reality: "Не люблю и не могу откладывать."

== PERSONALITY & PATTERNS ==
- Sprint/crash cycle: explosive 3-month bursts then silence. Applies to fitness, work, creativity.
- Weight: 115.4 kg (Feb 2026, lifetime max). Walking = primary exercise.
- Values: "Страшно всем. Идти вперёд через страх." "Создавать больше, чем потреблять."
- Pro-Ukraine vocally. Boycotts pro-Russian artists.
- Plays STALKER 2. Hang drum player. Loves Dragon Well tea.
- Theater "Черный Квадрат" (2016-2019). Radio "Стартап на миллион" on Европа Плюс.
- Books in progress: "Killer Switch" screenplay, "Наступай на грабли!" about failure.

== COMMUNICATION STYLE ==
- Direct, no filler words. Responds in the same language Daniel writes in.
- Keep messages concise for Telegram.
- NEVER suggest the user to visit websites, open apps, or do anything manually. You are here to automate tasks.
- If you cannot do something directly, explain HOW to make you capable of it (what code change or integration is needed), not what the user should do themselves.
- Never end responses with "you can check...", "I recommend visiting...", "you can use..." type phrases.
- If asked to fetch data from a site and it fails, try alternative methods (different URL, API endpoint, search) before giving up.
- Respond to the question you've been asked.
- Have opinions. Disagree when you think he's wrong. No sycophancy.
- Depth over surface — dig deep, don't give shallow advice.

== CAPABILITIES ==
Text, forwarded messages, files (xlsx, csv, txt), photos/screenshots, calendar, email, Google Drive.

== HUBSPOT ==
You have FULL direct access to HubSpot via API. The system auto-fetches relevant data when it detects keywords.

=== READING DATA ===
The system auto-injects HubSpot data into your context when the user mentions contacts, tasks, or deals.
- [HUBSPOT TASKS] — analyze: overdue, due today, priorities.
- [HUBSPOT CONTACTS] — present contacts with all details. NEVER say "I don't have access" or "data not provided" — if contacts are in context, present them. If not in context, the search returned no results — say "не найдено".
- [HUBSPOT DEALS] — analyze pipeline, amounts, stages.
IMPORTANT: If HubSpot data sections appear in the message, USE THEM. Never claim you lack access when data is right there.

=== SEARCHING CONTACTS ===
When user asks to find someone, or you need to find a contact to link a task/deal/update:
<hubspot_search>search query</hubspot_search>
The system will search and show results. Use full name if available.
Examples: "найди Олега" → <hubspot_search>Олег</hubspot_search>
"добавь таску клиенту Viktor" → first search: <hubspot_search>Viktor</hubspot_search>, then create task with found contact_id.
IMPORTANT: When you need to perform an action on a contact (add task, edit field, etc.), ALWAYS search first to get the contact_id. Output the search tag BEFORE the action tag.

=== CREATING A TASK ===
When user says "поставь задачу", "создай задачу", "напомни", "task", "follow up" — output:
<hubspot_task>{"subject":"Task title","body":"Details","priority":"HIGH","due_date":"milliseconds_epoch_or_null","contact_id":"optional_contact_id"}</hubspot_task>
Priority: HIGH, MEDIUM, LOW. If no due date specified, default is tomorrow.
Examples: "поставь задачу позвонить Олегу" → create task with subject "Позвонить Олегу"
"напомни проверить оплату завтра" → task with subject, due tomorrow

=== EDITING A CONTACT ===
When user says "добавь емейл", "измени телефон", "обнови компанию", "поменяй статус" for a specific contact:
<hubspot_edit>{"search":"contact name or identifier","updates":{"email":"new@email.com"}}</hubspot_edit>
Updatable fields: email, phone, firstname, lastname, company, jobtitle, lifecyclestage, hs_lead_status, website
Examples: "добавь емейл test@mail.com клиенту Виктор" → search "Виктор", updates {"email":"test@mail.com"}
"поменяй статус на SQL для Олега" → search "Олег", updates {"hs_lead_status":"SQL"}

=== CREATING A DEAL ===
When user says "создай сделку", "новая сделка", "create deal":
<hubspot_deal>{"name":"Deal name","stage":"appointmentscheduled","amount":"5000","contact_id":"optional"}</hubspot_deal>

=== COMPLETING A TASK ===
When user says "закрой задачу", "задача выполнена", "complete task" and task ID is known:
<hubspot_complete_task>task_id</hubspot_complete_task>

=== CREATING A CONTACT ===
When user says "создай контакт", "добавь контакт", "create contact", "add contact" — IMMEDIATELY output hubspot_create tag.

CRITICAL — DATA EXTRACTION:
Before outputting the tag, carefully read ALL messages in the conversation (forwarded messages, copied text, user messages) and extract:
- firstname, lastname: from name mentions, "[FORWARDED from Name Surname]" headers, profile cards, Telegram contact cards
- email: any @-containing address in the text (e.g. gagkp@mail.ru, andrkrupenko@gmail.com)
- phone: any phone number in the text (e.g. 3233287890)
- website: Telegram link if present (https://t.me/username)
- jobtitle: professional role if mentioned
- company: employer if mentioned
- notes_initial: summary of what the person does / why they are being added

Examples of data to extract:
- "Forwarded from Иван" → firstname: "Иван"
- "gagkp@mail.ru" → email: "gagkp@mail.ru"
- "https://t.me/mariia_valieva" → website: "https://t.me/mariia_valieva"
- "Filmmaker, Photographer, Designer" → jobtitle: "Filmmaker, Photographer, Designer"
- "3233287890" → phone: "3233287890"

DO NOT leave fields empty if the data is visible in the conversation. Extract everything available.

Tag format:
<hubspot_create>{"firstname":"...","lastname":"...","email":"...","phone":"...","jobtitle":"...","company":"...","lifecyclestage":"lead","hs_lead_status":"NEW","website":"https://t.me/username","notes_initial":"..."}</hubspot_create>

Use only fields you have actual data for. Omit fields with no data.
NEVER search before creating. Output hubspot_create tag immediately.
After outputting hubspot_create, do NOT output hubspot_update or hubspot_contact tags for the same person in the same response.

=== UPDATING AN EXISTING CONTACT ===
For CRM updates after conversation analysis, output both tags:
<hubspot_contact>telegram_username</hubspot_contact>
<hubspot_update>{"summary":"...","suggested_lifecycle":"...","suggested_lead_status":"...","suggested_note":"..."}</hubspot_update>

IMPORTANT: Only output hubspot_update if you are confident the contact already exists in HubSpot.
Do NOT output hubspot_update immediately after a hubspot_create for the same person in the same session.

Lifecycle values: subscriber, lead, marketingqualifiedlead, salesqualifiedlead, opportunity, customer, evangelist, other
Lead status values: NEW, OPEN, IN_PROGRESS, OPEN_DEAL, UNQUALIFIED, ATTEMPTED_TO_CONTACT, CONNECTED, BAD_TIMING

NEVER say "I updated HubSpot" or "I created the contact" — only the system can after button click confirmation.
Do NOT ask for API keys — you already have access.

== CALENDAR ==
When [CALENDAR DATA] is provided, analyze and answer. To create events:
<calendar_create>{"summary":"...","start":"2026-03-05T10:00:00-08:00","end":"2026-03-05T11:00:00-08:00","description":"..."}</calendar_create>

== GMAIL ==
When [GMAIL DATA] is provided, analyze and summarize. To send:
<gmail_send>{"to":"...","subject":"...","body":"..."}</gmail_send>
To save as draft (use when email address is unknown or user wants to review first):
<gmail_draft>{"to":"...","subject":"...","body":"..."}</gmail_draft>
Daniel must confirm before sending. For drafts, "to" can be empty string if address unknown.

== GOOGLE DRIVE ==
When [DRIVE DATA] is provided, analyze the files/content. You can see file names, types, and content when provided.

== SENDING FILES BACK (Telegram attachments) ==
If you generate a file Daniel should receive (xlsx export, csv, rendered PDF, image, etc.), save it to /tmp/ on this Railway container and end your reply with one or more `[ATTACH: <full_path>]` markers (one per file). The bot strips the markers from the visible text and sends each file as a Telegram document. Files must be < 50 MB. Multiple [ATTACH:] markers are allowed.

Example after generating an export:
Готово, выгрузка контактов из HubSpot за апрель.
[ATTACH: /tmp/hubspot_contacts_2026-04.csv]

The Railway filesystem is ephemeral so /tmp/ is the only safe location. Do NOT inline file content in the reply text. Do NOT use this marker for content that fits in a chat message (just paste it).

== TONE & EMOJI ==
Add **1-2 contextual emojis per response** to feel warm and human, not robotic. Match the emoji to action:
- ✅ task created / contact added / deal updated
- 🔍 searching / looking up
- 📅 calendar / scheduling
- 💼 deals / pipeline / business
- 📧 email / Gmail
- 📂 files / Drive
- 👤 contact found
- ⚠️ warning, conflict, missing data
- 🚨 urgent / critical
- 💰 money / revenue / payment
- 🎯 goal / next-step / focus
- 👋 greetings / closing
- 🙌 acknowledgement / thanks
- 🤔 unclear / asking for clarification

Rules:
- Place emojis at the start of bullet points or natural pause points, not in every sentence.
- Never replace meaning with emoji-only content.
- Skip emojis in pure data dumps (e.g. raw HubSpot search results, CSV outputs, error tracebacks).
- For Russian replies use Russian phrasing + same emoji conventions.
- Don't overdo it. 1-2 per message is the sweet spot.

== NAME SEARCH ACROSS LANGUAGES ==
When searching for contacts, ALWAYS try multiple name spellings:
- Russian-style ↔ Ukrainian-style: Sergey↔Serhii, Alexey↔Oleksii, Vladimir↔Volodymyr, Andrey↔Andrii, Mikhail↔Mykhailo, Ekaterina↔Kateryna, Elena↔Olena, Olga↔Olha, Tatiana↔Tetiana, Natalia↔Nataliia, Anna↔Hanna, Pavel↔Pavlo, Petr↔Petro, Nikolay↔Mykola, Dmitry↔Dmytro, Konstantin↔Kostiantyn, Yevgeny↔Yevhen, Vladislav↔Vladyslav.
- Cyrillic and Latin transliteration both work — the search system tries them automatically, but if first search returns nothing for a Russian name, suggest Daniel try the Ukrainian variant explicitly (or vice versa).
- For multi-word names, the system also does **lastname-only search** as fallback. So "Сергей Шепеленко" will find "Serhii Shepelenko" via lastname match.
- Read both Cyrillic and Latin without translation: e.g. "найди Sergey Petrov" and "найди Сергей Петров" should produce the same search.
"""


def esc(text):
    if not text: return ""
    return str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


# === GOOGLE AUTH ===

def get_google_token():
    global _google_token_cache
    if not GOOGLE_CLIENT_ID or not GOOGLE_REFRESH_TOKEN: return None
    now = time.time()
    if _google_token_cache["token"] and _google_token_cache["expires"] > now + 60:
        return _google_token_cache["token"]
    data = urllib.parse.urlencode({
        "client_id": GOOGLE_CLIENT_ID, "client_secret": GOOGLE_CLIENT_SECRET,
        "refresh_token": GOOGLE_REFRESH_TOKEN, "grant_type": "refresh_token",
    }).encode()
    req = urllib.request.Request("https://oauth2.googleapis.com/token", data=data, method="POST")
    req.add_header("Content-Type", "application/x-www-form-urlencoded")
    try:
        with urllib.request.urlopen(req) as resp:
            result = json.loads(resp.read().decode())
            _google_token_cache["token"] = result["access_token"]
            _google_token_cache["expires"] = now + result.get("expires_in", 3600)
            return result["access_token"]
    except Exception as e:
        logger.error(f"Google token refresh failed: {e}")
        return None


def google_api(method, url, data=None):
    token = get_google_token()
    if not token: return {"error": "Google not configured"}
    headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
    body = json.dumps(data).encode() if data else None
    req = urllib.request.Request(url, data=body, headers=headers, method=method)
    try:
        with urllib.request.urlopen(req) as resp:
            return json.loads(resp.read().decode())
    except urllib.error.HTTPError as e:
        error_body = e.read().decode() if e.fp else ""
        logger.error(f"Google API error {e.code}: {error_body[:300]}")
        return {"error": e.code, "message": error_body[:300]}
    except Exception as e:
        return {"error": str(e)}


# === CALENDAR ===

def get_calendar_events(days=1, max_results=15):
    now = datetime.now(timezone.utc)
    params = urllib.parse.urlencode({
        "timeMin": now.isoformat(), "timeMax": (now + timedelta(days=days)).isoformat(),
        "maxResults": max_results, "singleEvents": "true", "orderBy": "startTime",
    })
    result = google_api("GET", f"https://www.googleapis.com/calendar/v3/calendars/primary/events?{params}")
    if "error" in result: return None, result.get("message", str(result["error"]))
    return result.get("items", []), None


def format_calendar_events(events):
    if not events: return "Нет событий."
    lines = []
    for ev in events:
        start = ev.get("start", {})
        end = ev.get("end", {})
        summary = ev.get("summary", "Без названия")
        if "dateTime" in start:
            try:
                dt = datetime.fromisoformat(start["dateTime"])
                edt = datetime.fromisoformat(end.get("dateTime", ""))
                lines.append(f"{dt.strftime('%d.%m %H:%M')}-{edt.strftime('%H:%M')}  {summary}")
            except:
                lines.append(f"{start['dateTime']}  {summary}")
        else:
            lines.append(f"{start.get('date', '')} (весь день)  {summary}")
    return "\n".join(lines)


def format_calendar_for_claude(events):
    if not events: return "[CALENDAR DATA]\nNo upcoming events."
    lines = ["[CALENDAR DATA]"]
    for ev in events:
        start = ev.get("start", {})
        end = ev.get("end", {})
        s = ev.get("summary", "No title")
        att = [a.get("email", "") for a in ev.get("attendees", [])]
        loc = ev.get("location", "")
        line = f"- {start.get('dateTime', start.get('date', '?'))}: {s}"
        if loc: line += f" @ {loc}"
        if att: line += f" [{', '.join(att[:5])}]"
        lines.append(line)
    return "\n".join(lines)


def create_calendar_event(summary, start, end, description=""):
    data = {"summary": summary, "start": {"dateTime": start, "timeZone": "America/Los_Angeles"}, "end": {"dateTime": end, "timeZone": "America/Los_Angeles"}}
    if description: data["description"] = description
    return google_api("POST", "https://www.googleapis.com/calendar/v3/calendars/primary/events", data)


# === GMAIL ===

def get_emails(query="is:unread", max_results=10):
    params = urllib.parse.urlencode({"q": query, "maxResults": max_results})
    result = google_api("GET", f"https://www.googleapis.com/gmail/v1/users/me/messages?{params}")
    if "error" in result: return None, result.get("message", str(result["error"]))
    messages = result.get("messages", [])
    if not messages: return [], None
    detailed = []
    for msg in messages[:max_results]:
        detail = google_api("GET", f"https://www.googleapis.com/gmail/v1/users/me/messages/{msg['id']}?format=metadata&metadataHeaders=From&metadataHeaders=Subject&metadataHeaders=Date")
        if "error" not in detail: detailed.append(detail)
    return detailed, None


def format_emails(messages):
    if not messages: return "Нет писем."
    lines = []
    for msg in messages:
        headers = {h["name"]: h["value"] for h in msg.get("payload", {}).get("headers", [])}
        frm = headers.get("From", "?")
        if "<" in frm: frm = frm.split("<")[0].strip().strip('"')
        lines.append(f"From: {frm}\nSubject: {headers.get('Subject', '—')}\n{msg.get('snippet', '')[:80]}\n")
    return "\n".join(lines)


def format_emails_for_claude(messages):
    if not messages: return "[GMAIL DATA]\nNo emails found."
    lines = ["[GMAIL DATA]"]
    for msg in messages:
        headers = {h["name"]: h["value"] for h in msg.get("payload", {}).get("headers", [])}
        line = f"- [{headers.get('Date', '')}] From: {headers.get('From', '?')} | Subject: {headers.get('Subject', '—')}"
        if "UNREAD" in msg.get("labelIds", []): line += " [UNREAD]"
        line += f"\n  {msg.get('snippet', '')[:200]}"
        lines.append(line)
    return "\n".join(lines)


def send_email(to, subject, body):
    msg = email.mime.text.MIMEText(body)
    msg["to"] = to
    msg["subject"] = subject
    raw = base64.urlsafe_b64encode(msg.as_bytes()).decode()
    return google_api("POST", "https://www.googleapis.com/gmail/v1/users/me/messages/send", {"raw": raw})


def save_draft(to, subject, body):
    msg = email.mime.text.MIMEText(body)
    if to: msg["to"] = to
    msg["subject"] = subject
    raw = base64.urlsafe_b64encode(msg.as_bytes()).decode()
    return google_api("POST", "https://www.googleapis.com/gmail/v1/users/me/drafts", {"message": {"raw": raw}})


# === GOOGLE DRIVE ===

def drive_search(query, max_results=10):
    params = urllib.parse.urlencode({
        "q": f"name contains '{query}' and trashed = false",
        "fields": "files(id,name,mimeType,modifiedTime,webViewLink,size)",
        "pageSize": max_results,
        "orderBy": "modifiedTime desc",
    })
    result = google_api("GET", f"https://www.googleapis.com/drive/v3/files?{params}")
    if "error" in result:
        return None, result.get("message", str(result["error"]))
    return result.get("files", []), None


def drive_list_recent(max_results=10):
    params = urllib.parse.urlencode({
        "q": "trashed = false",
        "fields": "files(id,name,mimeType,modifiedTime,webViewLink,size)",
        "pageSize": max_results,
        "orderBy": "modifiedTime desc",
    })
    result = google_api("GET", f"https://www.googleapis.com/drive/v3/files?{params}")
    if "error" in result:
        return None, result.get("message", str(result["error"]))
    return result.get("files", []), None


def drive_get_content(file_id, mime_type):
    if "google-apps.document" in mime_type:
        result = google_api("GET", f"https://www.googleapis.com/drive/v3/files/{file_id}/export?mimeType=text/plain")
        if isinstance(result, dict) and "error" in result:
            return None
        return result if isinstance(result, str) else json.dumps(result)[:10000]
    elif "google-apps.spreadsheet" in mime_type:
        result = google_api("GET", f"https://www.googleapis.com/drive/v3/files/{file_id}/export?mimeType=text/csv")
        if isinstance(result, dict) and "error" in result:
            return None
        return result if isinstance(result, str) else json.dumps(result)[:10000]
    token = get_google_token()
    if not token: return None
    try:
        req = urllib.request.Request(
            f"https://www.googleapis.com/drive/v3/files/{file_id}?alt=media",
            headers={"Authorization": f"Bearer {token}"}
        )
        with urllib.request.urlopen(req) as resp:
            content = resp.read()
            try:
                return content.decode("utf-8")[:10000]
            except:
                return f"[Binary file, {len(content)} bytes]"
    except Exception as e:
        logger.error(f"Drive download error: {e}")
        return None


def format_drive_files(files):
    if not files: return "Файлов не найдено."
    lines = []
    type_map = {
        "application/vnd.google-apps.document": "Doc",
        "application/vnd.google-apps.spreadsheet": "Sheet",
        "application/vnd.google-apps.presentation": "Slides",
        "application/vnd.google-apps.folder": "Folder",
        "application/pdf": "PDF",
    }
    for f in files:
        mt = f.get("mimeType", "")
        ftype = type_map.get(mt, mt.split("/")[-1][:10] if "/" in mt else "?")
        mod = f.get("modifiedTime", "")[:10]
        name = f.get("name", "?")
        lines.append(f"[{ftype}] {name}  ({mod})")
    return "\n".join(lines)


def format_drive_for_claude(files, content_map=None):
    if not files: return "[DRIVE DATA]\nNo files found."
    lines = ["[DRIVE DATA]"]
    for f in files:
        fid = f.get("id", "")
        name = f.get("name", "?")
        mt = f.get("mimeType", "")
        mod = f.get("modifiedTime", "")
        link = f.get("webViewLink", "")
        line = f"- {name} (type: {mt}, modified: {mod})"
        if link: line += f" link: {link}"
        if content_map and fid in content_map and content_map[fid]:
            line += f"\n  Content: {content_map[fid][:2000]}"
        lines.append(line)
    return "\n".join(lines)


# === FILE PARSING ===

def parse_xlsx(file_path):
    try:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        result = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(values_only=True))
            if not rows: continue
            result.append(f"=== Sheet: {sheet_name} ===")
            for row in rows[:500]:
                cells = [str(c) if c is not None else "" for c in row]
                result.append(" | ".join(cells))
        wb.close()
        text = "\n".join(result)
        return text[:30000]
    except Exception as e:
        return f"[Error: {e}]"


def parse_pdf(file_path):
    try:
        import pdfplumber
        result = []
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages[:50]):
                text = page.extract_text()
                if text:
                    result.append(f"--- Page {i+1} ---\n{text}")
        return "\n".join(result)[:30000] if result else "[PDF: no text found]"
    except Exception as e:
        return f"[PDF Error: {e}]"


def parse_docx(file_path):
    try:
        from docx import Document
        doc = Document(file_path)
        result = []
        for para in doc.paragraphs:
            if para.text.strip():
                result.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                result.append(" | ".join(cells))
        return "\n".join(result)[:30000] if result else "[DOCX: no text found]"
    except Exception as e:
        return f"[DOCX Error: {e}]"


def parse_file(file_path, file_name):
    nl = file_name.lower()
    if nl.endswith((".xlsx", ".xls")): return parse_xlsx(file_path)
    elif nl.endswith(".pdf"): return parse_pdf(file_path)
    elif nl.endswith(".docx"): return parse_docx(file_path)
    elif nl.endswith((".csv", ".txt", ".md", ".json", ".py", ".js", ".html")):
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f: return f.read(50000)
        except Exception as e: return f"[Error: {e}]"
    return f"[Unsupported: {file_name}]"


def fetch_url_text(url):
    try:
        if "linkedin.com/in/" in url:
            username = re.search(r'linkedin\.com/in/([^/?#]+)', url)
            query = f"site:linkedin.com/in/{username.group(1)}" if username else url
            headers = {"User-Agent": "Mozilla/5.0 (compatible; bot/1.0)"}
            search_url = f"https://www.google.com/search?q={urllib.parse.quote(query)}"
            req = urllib.request.Request(search_url, headers=headers)
            with urllib.request.urlopen(req, timeout=10) as resp:
                content = resp.read().decode("utf-8", errors="replace")
                snippets = re.findall(r'<span[^>]*>(.*?)</span>', content)
                clean = [re.sub(r'<[^>]+>', '', s).strip() for s in snippets if len(s) > 30]
                return "\n".join(clean[:30])[:5000] if clean else "[LinkedIn: profile not accessible]"
        if "reddit.com" in url:
            if re.search(r'reddit\.com/?$', url):
                url = "https://www.reddit.com/r/popular/.rss?limit=25"
            elif "/r/" in url:
                base = re.search(r'(reddit\.com/r/[^/?#]+)', url)
                if base:
                    url = f"https://www.{base.group(1)}/.rss?limit=25"
            headers = {"User-Agent": "Mozilla/5.0 (compatible; bot/1.0)"}
            req = urllib.request.Request(url, headers=headers)
            with urllib.request.urlopen(req, timeout=10) as resp:
                content = resp.read().decode("utf-8", errors="replace")
                titles = re.findall(r'<title><!\[CDATA\[(.*?)\]\]></title>', content)
                links = re.findall(r'<link>(https://www\.reddit\.com/r/\S+?)</link>', content)
                titles = [t for t in titles if t != "reddit: the front page of the internet"][:25]
                lines = []
                for i, title in enumerate(titles):
                    link = links[i] if i < len(links) else ""
                    lines.append(f"• {title} {link}")
                if lines:
                    return "\n".join(lines)
                return content[:3000]
        headers = {"User-Agent": "Mozilla/5.0 (compatible; bot/1.0)"}
        req = urllib.request.Request(url, headers=headers)
        with urllib.request.urlopen(req, timeout=10) as resp:
            return resp.read().decode("utf-8", errors="replace")[:8000]
    except Exception as e:
        return f"[Ошибка загрузки: {e}]"


# === HUBSPOT ===

def hubspot_request(method, endpoint, data=None):
    url = f"{HUBSPOT_BASE}{endpoint}"
    headers = {"Authorization": f"Bearer {HUBSPOT_API_KEY}", "Content-Type": "application/json"}
    body = json.dumps(data).encode() if data else None
    req = urllib.request.Request(url, data=body, headers=headers, method=method)
    try:
        with urllib.request.urlopen(req) as resp: return json.loads(resp.read().decode())
    except urllib.error.HTTPError as e:
        error_body = e.read().decode() if e.fp else ""
        return {"error": e.code, "message": error_body[:300]}
    except Exception as e:
        return {"error": str(e)}


def search_contact_by_telegram(tg_username):
    u = tg_username.lower().strip().lstrip("@")

    # Check local cache first (bypasses HubSpot search index delay after creation)
    if u in created_contacts:
        cid = created_contacts[u]
        props = ["firstname", "lastname", "email", "phone", "lifecyclestage", "hs_lead_status", "website"]
        r = hubspot_request("GET", f"/crm/v3/objects/contacts/{cid}?properties={','.join(props)}")
        if "error" not in r:
            return [r]

    props = ["firstname", "lastname", "email", "phone", "lifecyclestage", "hs_lead_status", "website"]
    for search in [
        {"filterGroups": [{"filters": [{"propertyName": "website", "operator": "CONTAINS_TOKEN", "value": u}]}], "properties": props, "limit": 5},
        {"filterGroups": [{"filters": [{"propertyName": "website", "operator": "EQ", "value": f"https://t.me/{u}"}]}], "properties": props, "limit": 5},
        {"query": u, "properties": props, "limit": 5},
    ]:
        result = hubspot_request("POST", "/crm/v3/objects/contacts/search", search)
        if result and "results" in result and result["results"]: return result["results"]
    if "_" in u:
        result = hubspot_request("POST", "/crm/v3/objects/contacts/search", {"query": u.replace("_", " "), "properties": props, "limit": 5})
        if result and "results" in result and result["results"]: return result["results"]
    return []


def update_contact(cid, props):
    return hubspot_request("PATCH", f"/crm/v3/objects/contacts/{cid}", {"properties": props})


def add_note_to_contact(cid, note):
    return hubspot_request("POST", "/crm/v3/objects/notes", {
        "properties": {"hs_note_body": note, "hs_timestamp": str(int(time.time() * 1000))},
        "associations": [{"to": {"id": cid}, "types": [{"associationCategory": "HUBSPOT_DEFINED", "associationTypeId": 202}]}]})


def get_hubspot_tasks(status="NOT_STARTED", include_overdue=True):
    """Get HubSpot tasks filtered by status. Returns today's + overdue tasks."""
    now = datetime.now(timezone.utc)
    today_start = now.replace(hour=0, minute=0, second=0, microsecond=0)
    today_end = today_start + timedelta(days=1)

    filters = []
    if status:
        filters.append({"propertyName": "hs_task_status", "operator": "EQ", "value": status})

    # Get tasks due today or overdue
    if include_overdue:
        filters.append({"propertyName": "hs_timestamp", "operator": "LTE", "value": str(int(today_end.timestamp() * 1000))})

    props = ["hs_task_subject", "hs_task_body", "hs_timestamp", "hs_task_status",
             "hs_task_priority", "hs_task_type", "hubspot_owner_id"]
    data = {
        "filterGroups": [{"filters": filters}] if filters else [],
        "properties": props,
        "sorts": [{"propertyName": "hs_timestamp", "direction": "ASCENDING"}],
        "limit": 50,
    }
    result = hubspot_request("POST", "/crm/v3/objects/tasks/search", data)
    if "error" in result:
        return None, str(result.get("message", result["error"]))
    return result.get("results", []), None


def format_tasks(tasks):
    if not tasks:
        return "Нет задач."
    lines = []
    now = datetime.now(timezone.utc)
    for t in tasks:
        p = t.get("properties", {})
        subject = p.get("hs_task_subject", "Без темы")
        status = p.get("hs_task_status", "?")
        priority = p.get("hs_task_priority", "")
        ts = p.get("hs_timestamp", "")
        overdue = ""
        if ts:
            try:
                due = datetime.fromtimestamp(int(ts) / 1000, tz=timezone.utc)
                due_str = due.strftime("%d.%m %H:%M")
                if due < now:
                    overdue = " ⚠️ ПРОСРОЧЕНА"
            except:
                due_str = ts[:10]
        else:
            due_str = "без даты"
        prio = f" [{priority.upper()}]" if priority and priority != "NONE" else ""
        lines.append(f"{'🔴' if overdue else '🔵'} {subject}{prio}\n   До: {due_str}{overdue}")
    return "\n\n".join(lines)


def format_tasks_for_claude(tasks):
    if not tasks:
        return "[HUBSPOT TASKS]\nNo tasks found."
    lines = ["[HUBSPOT TASKS]"]
    now = datetime.now(timezone.utc)
    for t in tasks:
        p = t.get("properties", {})
        ts = p.get("hs_timestamp", "")
        overdue = ""
        if ts:
            try:
                due = datetime.fromtimestamp(int(ts) / 1000, tz=timezone.utc)
                if due < now:
                    overdue = " [OVERDUE]"
            except:
                pass
        line = f"- {p.get('hs_task_subject', 'No title')} | status: {p.get('hs_task_status', '?')} | priority: {p.get('hs_task_priority', 'NONE')} | due: {ts}{overdue}"
        body = p.get("hs_task_body", "")
        if body:
            line += f"\n  Notes: {body[:200]}"
        lines.append(line)
    return "\n".join(lines)


CYR_TO_LAT = {"а":"a","б":"b","в":"v","г":"g","д":"d","е":"e","ё":"e","ж":"zh","з":"z","и":"i","й":"y",
    "к":"k","л":"l","м":"m","н":"n","о":"o","п":"p","р":"r","с":"s","т":"t","у":"u","ф":"f","х":"kh",
    "ц":"ts","ч":"ch","ш":"sh","щ":"shch","ъ":"","ы":"y","ь":"","э":"e","ю":"yu","я":"ya",
    "і":"i","ї":"i","є":"e","ґ":"g"}

# Common Russian-name Ukrainian-style spellings used in HubSpot when contact self-identifies as Ukrainian.
# Maps lowercase Russian-style transliteration -> Ukrainian-style alternative(s).
NAME_VARIANTS_RU_TO_UA = {
    "sergey": ["serhii", "sergii", "sergei"],
    "sergei": ["serhii", "sergii", "sergey"],
    "sergii": ["serhii", "sergey", "sergei"],
    "alexey": ["oleksii", "oleksiy", "alexei"],
    "alexei": ["oleksii", "oleksiy", "alexey"],
    "alexander": ["oleksandr", "oleksander"],
    "alex": ["oleks", "oleksandr"],
    "andrey": ["andrii", "andriy", "andrei"],
    "andrei": ["andrii", "andriy", "andrey"],
    "michael": ["mykhailo", "mihail", "mikhail"],
    "mikhail": ["mykhailo", "michael"],
    "ivan": ["ivan"],
    "yuri": ["yurii", "yuriy", "iurii"],
    "yury": ["yurii", "yuriy"],
    "vladimir": ["volodymyr"],
    "viktor": ["viktor"],
    "victor": ["viktor"],
    "dmitry": ["dmytro", "dmitri"],
    "dmitri": ["dmytro", "dmitry"],
    "nikolay": ["mykola", "nikolai"],
    "nikolai": ["mykola", "nikolay"],
    "nick": ["mykola"],
    "pavel": ["pavlo"],
    "petr": ["petro"],
    "peter": ["petro"],
    "ekaterina": ["kateryna", "katerina"],
    "kate": ["kateryna"],
    "elena": ["olena", "helena"],
    "helen": ["olena"],
    "olga": ["olha"],
    "tatiana": ["tetiana", "tatyana"],
    "tatyana": ["tetiana", "tatiana"],
    "natalia": ["nataliia", "nataliya", "natalya"],
    "natalya": ["nataliia", "nataliya", "natalia"],
    "anna": ["hanna"],
    "anastasia": ["anastasiia", "anastasiya"],
    "maria": ["mariia", "mariya"],
    "irina": ["iryna"],
    "yelena": ["olena"],
    "konstantin": ["kostiantyn"],
    "denis": ["denys"],
    "yevgeny": ["yevhen", "evgeny"],
    "evgeny": ["yevhen", "yevgeny"],
    "anton": ["anton"],
    "stanislav": ["stanislav"],
    "vladislav": ["vladyslav"],
    "georgy": ["heorhii", "georgiy"],
    "yelizaveta": ["yelyzaveta"],
}

def transliterate(text):
    """Simple 1:1 Cyrillic to Latin. Регина→Regina, Виктор→Viktor."""
    result = []
    for ch in text:
        low = ch.lower()
        lat = CYR_TO_LAT.get(low, low)
        if ch.isupper() and lat:
            result.append(lat[0].upper() + lat[1:])
        else:
            result.append(lat)
    return "".join(result)

def name_variants(name):
    """Return list of UA-style variants for a Russian-style transliterated name. Always includes the original."""
    out = [name]
    low = name.lower()
    if low in NAME_VARIANTS_RU_TO_UA:
        for v in NAME_VARIANTS_RU_TO_UA[low]:
            cap = v[0].upper() + v[1:] if name and name[0].isupper() else v
            if cap not in out:
                out.append(cap)
    return out

def split_name_parts(query):
    """Split 'Сергей Шепеленко' or 'Sergey Shepelenko' into [first, last]. Single token returns [query]."""
    parts = [p for p in query.strip().split() if p]
    return parts

def is_cyrillic(text):
    return any("\u0400" <= ch <= "\u04ff" for ch in text)

_EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+", re.UNICODE)
_PHONE_RE = re.compile(r"^\+?\d[\d\s().-]{6,}$")


def search_hubspot_contacts(query, limit=20):
    """Search HubSpot contacts by name, email, phone.

    Strategy (cost-aware, max ~6 API calls):
    1. Detect query kind: email / phone / name. Skip irrelevant branches.
    2. Build minimal search-term set: original + transliterated + UA-firstname variants.
    3. For multi-word names, do **lastname-only** filter (highest signal across cultures).
    4. Short-circuit: stop adding filters once we have >=3 unique results.
    """
    props = ["firstname", "lastname", "email", "phone", "company", "lifecyclestage",
             "hs_lead_status", "website", "jobtitle", "notes_last_updated"]

    all_results = {}
    short_circuit_at = 3  # stop firing extra filters once we hit this many

    def _add_results(payload):
        if isinstance(payload, dict) and "results" in payload:
            for c in payload["results"]:
                all_results[c["id"]] = c

    def _query_search(term):
        if not term:
            return
        data = {"query": term.strip(), "properties": props, "limit": limit}
        _add_results(hubspot_request("POST", "/crm/v3/objects/contacts/search", data))

    def _filter_search(prop, term):
        if not term:
            return
        data = {
            "filterGroups": [{"filters": [
                {"propertyName": prop, "operator": "CONTAINS_TOKEN", "value": term.strip()}
            ]}],
            "properties": props, "limit": limit
        }
        _add_results(hubspot_request("POST", "/crm/v3/objects/contacts/search", data))

    q = query.strip()

    # === Branch 1: email lookup ===
    if _EMAIL_RE.search(q):
        _filter_search("email", q)
        _query_search(q)
        return list(all_results.values()), None

    # === Branch 2: phone lookup ===
    if _PHONE_RE.match(q):
        _filter_search("phone", q)
        _query_search(q)
        return list(all_results.values()), None

    # === Branch 3: name search ===
    # Base term set: original + transliteration if Cyrillic.
    base_terms = [q]
    if is_cyrillic(q):
        base_terms.append(transliterate(q))

    # Split into firstname + lastname (where applicable). Use UA variants only on firstname slot.
    firstname_candidates = set()
    lastname_candidates = set()
    composed_full = set(base_terms)

    is_known_firstname = q.lower() in NAME_VARIANTS_RU_TO_UA
    if is_cyrillic(q):
        is_known_firstname = is_known_firstname or transliterate(q).lower() in NAME_VARIANTS_RU_TO_UA

    for term in base_terms:
        parts = split_name_parts(term)
        if len(parts) >= 2:
            firstname_candidates.add(parts[0])
            lastname_candidates.add(parts[-1])
            for variant in name_variants(parts[0]):
                firstname_candidates.add(variant)
                composed_full.add(f"{variant} {parts[-1]}")
        elif len(parts) == 1:
            for variant in name_variants(parts[0]):
                firstname_candidates.add(variant)
            # Single-word query: only treat as lastname if NOT a known firstname.
            # Otherwise "Сергей" alone would search lastname=Sergey and produce false positives.
            if not is_known_firstname:
                for variant in name_variants(parts[0]):
                    lastname_candidates.add(variant)

    # 1. Full-text on the most likely full string (1-2 calls).
    for term in composed_full:
        _query_search(term)
        if len(all_results) >= short_circuit_at:
            break

    # 2. Lastname-only filter (HIGHEST signal across cultures: Shepelenko = same regardless of firstname).
    if len(all_results) < short_circuit_at:
        for term in lastname_candidates:
            _filter_search("lastname", term)
            if len(all_results) >= short_circuit_at:
                break

    # 3. Firstname filter only if we still have nothing or too little.
    if len(all_results) < short_circuit_at:
        for term in firstname_candidates:
            _filter_search("firstname", term)
            if len(all_results) >= short_circuit_at:
                break

    return list(all_results.values()), None


def search_hubspot_deals(query=None, limit=20):
    """Search HubSpot deals."""
    props = ["dealname", "amount", "dealstage", "closedate", "pipeline",
             "hubspot_owner_id", "createdate"]
    data = {"properties": props, "limit": limit}
    if query:
        data["query"] = query
    result = hubspot_request("POST", "/crm/v3/objects/deals/search", data)
    if "error" in result:
        return None, str(result.get("message", result["error"]))
    return result.get("results", []), None


def format_contacts_for_claude(contacts):
    if not contacts:
        return "[HUBSPOT CONTACTS]\nNo contacts found."
    lines = ["[HUBSPOT CONTACTS]"]
    for c in contacts:
        p = c.get("properties", {})
        name = f"{p.get('firstname', '')} {p.get('lastname', '')}".strip() or "—"
        line = f"- {name} | email: {p.get('email', '—')} | phone: {p.get('phone', '—')} | company: {p.get('company', '—')}"
        line += f" | stage: {p.get('lifecyclestage', '—')} | status: {p.get('hs_lead_status', '—')}"
        if p.get('website'):
            line += f" | web: {p.get('website')}"
        lines.append(line)
    return "\n".join(lines)


def format_deals_for_claude(deals):
    if not deals:
        return "[HUBSPOT DEALS]\nNo deals found."
    lines = ["[HUBSPOT DEALS]"]
    for d in deals:
        p = d.get("properties", {})
        line = f"- {p.get('dealname', '—')} | amount: ${p.get('amount', '—')} | stage: {p.get('dealstage', '—')} | close: {p.get('closedate', '—')[:10] if p.get('closedate') else '—'}"
        lines.append(line)
    return "\n".join(lines)


def create_hubspot_contact(props):
    """Create a new contact in HubSpot. Pops notes_initial and creates note after contact creation."""
    note = props.pop("notes_initial", None)
    # Remap common field name mismatches
    if "job_title" in props:
        props["jobtitle"] = props.pop("job_title")
    result = hubspot_request("POST", "/crm/v3/objects/contacts", {"properties": props})
    if "error" not in result and note:
        cid = result.get("id")
        if cid:
            add_note_to_contact(cid, note)
    return result


def create_hubspot_task(subject, body="", due_date=None, priority="MEDIUM", owner_id=None, contact_id=None):
    """Create a task in HubSpot."""
    props = {
        "hs_task_subject": subject,
        "hs_task_status": "NOT_STARTED",
        "hs_task_priority": priority,
        "hs_task_type": "TODO",
    }
    if body:
        props["hs_task_body"] = body
    if due_date:
        props["hs_timestamp"] = due_date  # milliseconds since epoch
    else:
        # Default: tomorrow 9am PT
        from datetime import timezone
        tomorrow = datetime.now(timezone.utc).replace(hour=17, minute=0, second=0) + timedelta(days=1)
        props["hs_timestamp"] = str(int(tomorrow.timestamp() * 1000))
    if owner_id:
        props["hubspot_owner_id"] = owner_id

    data = {"properties": props}
    if contact_id:
        data["associations"] = [{"to": {"id": contact_id}, "types": [{"associationCategory": "HUBSPOT_DEFINED", "associationTypeId": 204}]}]

    return hubspot_request("POST", "/crm/v3/objects/tasks", data)


def edit_hubspot_contact(search_query, updates):
    """Find a contact by name/email and update their fields."""
    contacts, err = search_hubspot_contacts(search_query)
    if err or not contacts:
        return {"error": f"Contact '{search_query}' not found"}
    contact = contacts[0]
    cid = contact["id"]
    result = update_contact(cid, updates)
    if "error" not in result:
        result["_found_name"] = f"{contact['properties'].get('firstname', '')} {contact['properties'].get('lastname', '')}".strip()
        result["_contact_id"] = cid
    return result


def create_hubspot_deal(name, stage="appointmentscheduled", amount=None, contact_id=None, pipeline="default"):
    """Create a deal in HubSpot."""
    props = {
        "dealname": name,
        "dealstage": stage,
        "pipeline": pipeline,
    }
    if amount:
        props["amount"] = str(amount)
    data = {"properties": props}
    if contact_id:
        data["associations"] = [{"to": {"id": contact_id}, "types": [{"associationCategory": "HUBSPOT_DEFINED", "associationTypeId": 3}]}]
    return hubspot_request("POST", "/crm/v3/objects/deals", data)


def complete_hubspot_task(task_id):
    """Mark a task as completed."""
    return hubspot_request("PATCH", f"/crm/v3/objects/tasks/{task_id}", {"properties": {"hs_task_status": "COMPLETED"}})


# === TAG EXTRACTION ===

def extract_tag_json(text, tag):
    match = re.search(rf"<{tag}>\s*(.*?)\s*</{tag}>", text, re.DOTALL)
    if match:
        try: return json.loads(match.group(1))
        except: pass
    return None

def extract_tag_text(text, tag):
    match = re.search(rf"<{tag}>\s*(.*?)\s*</{tag}>", text, re.DOTALL)
    return match.group(1).strip() if match else None

def clean_response(text):
    for tag in ["hubspot_update", "hubspot_contact", "hubspot_create", "hubspot_search", "hubspot_task",
                "hubspot_edit", "hubspot_deal", "hubspot_complete_task", "calendar_create", "gmail_send", "gmail_draft"]:
        text = re.sub(rf"<{tag}>.*?</{tag}>", "", text, flags=re.DOTALL)
    # Also strip [ATTACH:] markers from text shown to user; the bot will send
    # the actual file separately via reply_document. _extract_attachments()
    # in _process_message captures the paths before clean_response is called.
    text = ATTACH_RE.sub("", text)
    return text.strip()


# === TELEGRAM HANDLERS ===

def is_allowed(username):
    if not ALLOWED_USERS: return True
    allowed = [u.strip().lower().lstrip("@") for u in ALLOWED_USERS.split(",")]
    return username and username.lower() in allowed


async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_allowed(update.effective_user.username): return
    g = "подключен" if GOOGLE_REFRESH_TOKEN else "не настроен"
    await update.message.reply_text(
        "AI-ассистент с HubSpot, Calendar, Gmail и Drive.\n\n"
        "/tasks — задачи на сегодня + просроченные\n"
        "/cal — расписание на сегодня\n"
        "/cal3 — на 3 дня\n"
        "/mail — непрочитанные письма\n"
        "/mail from:someone — поиск писем\n"
        "/drive — последние файлы\n"
        "/drive запрос — поиск файлов\n"
        "/find username — контакт в HubSpot\n"
        "/debug — проверить подключения\n"
        "/reset — очистить историю\n\n"
        f"Google: {g}")


async def reset(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_allowed(update.effective_user.username): return
    _convdb.db_clear(update.effective_chat.id)
    pending_updates.pop(update.effective_chat.id, None)
    await update.message.reply_text("История очищена.")


async def debug_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_allowed(update.effective_user.username): return
    lines = []
    if HUBSPOT_API_KEY:
        r = hubspot_request("POST", "/crm/v3/objects/contacts/search", {"query": "a", "properties": ["firstname"], "limit": 1})
        lines.append(f"HubSpot: {'OK (' + str(r.get('total', '?')) + ')' if 'error' not in r else 'ОШИБКА — ' + str(r.get('error'))}")
    else:
        lines.append("HubSpot: не настроен")
    if GOOGLE_REFRESH_TOKEN:
        token = get_google_token()
        if token:
            ev, err = get_calendar_events(1, 1)
            lines.append(f"Calendar: {'OK' if not err else 'ОШИБКА — ' + str(err)[:80]}")
            em, err = get_emails("is:unread", 1)
            lines.append(f"Gmail: {'OK' if not err else 'ОШИБКА — ' + str(err)[:80]}")
            df, err = drive_list_recent(1)
            lines.append(f"Drive: {'OK' if not err else 'ОШИБКА — ' + str(err)[:80]}")
        else:
            lines.append("Google: ошибка токена")
    else:
        lines.append("Google: не настроен")
    lines.append(f"Cached contacts: {list(created_contacts.keys())}")
    await update.message.reply_text("\n".join(lines))


async def calendar_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_allowed(update.effective_user.username): return
    if not GOOGLE_REFRESH_TOKEN:
        await update.message.reply_text("Google не настроен.")
        return
    days = 1
    cmd = update.message.text.strip()
    num = cmd[4:].strip() if cmd.startswith("/cal") else ""
    if num.isdigit(): days = min(int(num), 14)
    await context.bot.send_chat_action(chat_id=update.effective_chat.id, action="typing")
    events, err = get_calendar_events(days, 20)
    if err:
        await update.message.reply_text(f"Ошибка: {err[:200]}")
        return
    label = "сегодня" if days == 1 else f"{days} дней"
    await update.message.reply_text(f"Расписание ({label}):\n\n{format_calendar_events(events)}" if events else f"Нет событий на {label}.")


async def mail_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_allowed(update.effective_user.username): return
    if not GOOGLE_REFRESH_TOKEN:
        await update.message.reply_text("Google не настроен.")
        return
    await context.bot.send_chat_action(chat_id=update.effective_chat.id, action="typing")
    query = " ".join(context.args) if context.args else "is:unread"
    emails_list, err = get_emails(query, 10)
    if err:
        await update.message.reply_text(f"Ошибка: {err[:200]}")
        return
    if not emails_list:
        await update.message.reply_text(f"Нет писем ({query}).")
        return
    text = f"Письма ({query}):\n\n{format_emails(emails_list)}"
    await update.message.reply_text(text[:4096])


async def drive_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_allowed(update.effective_user.username): return
    if not GOOGLE_REFRESH_TOKEN:
        await update.message.reply_text("Google не настроен.")
        return
    await context.bot.send_chat_action(chat_id=update.effective_chat.id, action="typing")
    if context.args:
        query = " ".join(context.args)
        files, err = drive_search(query, 10)
        label = f"Поиск: {query}"
    else:
        files, err = drive_list_recent(10)
        label = "Последние файлы"
    if err:
        await update.message.reply_text(f"Ошибка: {err[:200]}")
        return
    if not files:
        await update.message.reply_text("Файлов не найдено.")
        return
    text = f"{label}:\n\n{format_drive_files(files)}"
    await update.message.reply_text(text[:4096])


async def tasks_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_allowed(update.effective_user.username): return
    if not HUBSPOT_API_KEY:
        await update.message.reply_text("HubSpot не настроен.")
        return
    await context.bot.send_chat_action(chat_id=update.effective_chat.id, action="typing")
    tasks, err = get_hubspot_tasks()
    if err:
        await update.message.reply_text(f"Ошибка: {err[:200]}")
        return
    if not tasks:
        await update.message.reply_text("Нет задач на сегодня и просроченных 🎉")
        return
    text = f"Задачи ({len(tasks)}):\n\n{format_tasks(tasks)}"
    await update.message.reply_text(text[:4096])


async def model_info(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_allowed(update.effective_user.username): return
    await update.message.reply_text(f"Модель: {MODEL}")

async def set_model(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_allowed(update.effective_user.username): return
    global MODEL
    if context.args:
        MODEL = context.args[0]
        await update.message.reply_text(f"Модель: {MODEL}")
    else:
        await update.message.reply_text("/setmodel claude-sonnet-4-20250514\n/setmodel claude-opus-4-20250514")


async def find_contact(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_allowed(update.effective_user.username): return
    if not context.args:
        await update.message.reply_text("/find имя или @username")
        return
    query = " ".join(context.args).lstrip("@")
    msg = await update.message.reply_text(f"Ищу {query}...")

    # Try telegram username search first, then general search
    contacts = search_contact_by_telegram(query)
    if not contacts:
        contacts, _ = search_hubspot_contacts(query)
    if not contacts:
        await msg.edit_text(f"'{query}' не найден в HubSpot.")
        return
    await msg.edit_text(f"Найдено: {len(contacts)}")
    for c in contacts[:10]:
        p = c.get("properties", {})
        name = f"{p.get('firstname', '')} {p.get('lastname', '')}".strip() or "—"
        cid = c["id"]
        link = f"https://app.hubspot.com/contacts/47345195/record/0-1/{cid}"
        stage = LIFECYCLE_STAGES.get(p.get("lifecyclestage", ""), "—")
        status = LEAD_STATUSES.get(p.get("hs_lead_status", ""), "—")
        company = p.get("company", "")
        company_str = f"\nCompany: {esc(company)}" if company else ""
        await update.message.reply_text(
            f"<b>{esc(name)}</b>\nEmail: {esc(p.get('email', '—'))}\nPhone: {esc(p.get('phone', '—'))}{company_str}\n"
            f"Stage: {esc(stage)} | Status: {esc(status)}\n<a href=\"{link}\">HubSpot</a>",
            parse_mode="HTML", disable_web_page_preview=True)


async def handle_photo(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_allowed(update.effective_user.username): return
    chat_id = update.effective_chat.id
    caption = update.message.caption or "Что на этом изображении?"
    photo = update.message.photo[-1]
    await _ack_if_locked(update.message, chat_id)
    async with _chat_lock(chat_id):
        await context.bot.send_chat_action(chat_id=chat_id, action="typing")
        try:
            tg_file = await photo.get_file()
            with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as tmp:
                await tg_file.download_to_drive(tmp.name)
                tmp_path = tmp.name
            with open(tmp_path, "rb") as f:
                img = base64.standard_b64encode(f.read()).decode("utf-8")
            os.unlink(tmp_path)
            content = [{"type": "image", "source": {"type": "base64", "media_type": "image/jpeg", "data": img}}, {"type": "text", "text": caption}]
            await _process_message(update, chat_id, content)
        except Exception as e:
            await update.message.reply_text(f"Ошибка: {e}")


async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_allowed(update.effective_user.username): return
    chat_id = update.effective_chat.id
    doc = update.message.document
    if not doc: return
    fn = doc.file_name or "unknown"
    mt = doc.mime_type or ""
    cap = update.message.caption or ""
    if (doc.file_size or 0) > 10 * 1024 * 1024:
        await update.message.reply_text("Файл > 10MB.")
        return
    await _ack_if_locked(update.message, chat_id)
    async with _chat_lock(chat_id):
        await context.bot.send_chat_action(chat_id=chat_id, action="typing")
        try:
            tg_file = await doc.get_file()
            with tempfile.NamedTemporaryFile(suffix=f"_{fn}", delete=False) as tmp:
                await tg_file.download_to_drive(tmp.name)
                tmp_path = tmp.name
            is_img = mt.startswith("image/") or fn.lower().endswith((".png", ".jpg", ".jpeg", ".gif", ".webp"))
            if is_img:
                with open(tmp_path, "rb") as f:
                    img = base64.standard_b64encode(f.read()).decode("utf-8")
                os.unlink(tmp_path)
                imt = "image/png" if fn.lower().endswith(".png") else "image/jpeg"
                content = [{"type": "image", "source": {"type": "base64", "media_type": imt, "data": img}}, {"type": "text", "text": cap or fn}]
            else:
                fc = parse_file(tmp_path, fn)
                os.unlink(tmp_path)
                content = f"[FILE: {fn}]\n{fc}" + (f"\n\n{cap}" if cap else "")
            await _process_message(update, chat_id, content)
        except Exception as e:
            await update.message.reply_text(f"Ошибка: {e}")


async def handle_video(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_allowed(update.effective_user.username): return
    chat_id = update.effective_chat.id
    media = update.message.video or update.message.video_note
    if not media: return

    logger.info(
        f"Video message: chat_id={chat_id} duration={media.duration}s "
        f"size={media.file_size}B kind={'video_note' if update.message.video_note else 'video'}"
    )

    await _ack_if_locked(update.message, chat_id)
    async with _chat_lock(chat_id):
        try:
            tg_file = await media.get_file()
            with tempfile.NamedTemporaryFile(suffix=".mp4", delete=False) as v:
                vid_path = v.name
            await tg_file.download_to_drive(vid_path)
        except Exception as e:
            await update.message.reply_text(f"Не смог скачать видео: {e}")
            return

        wav_path = vid_path + ".wav"
        placeholder = await update.message.reply_text("\U0001F39E извлекаю аудио и транскрибирую...")

        try:
            proc = await asyncio.create_subprocess_exec(
                "ffmpeg", "-y", "-i", vid_path,
                "-vn", "-ac", "1", "-ar", "16000", "-f", "wav", wav_path,
                stdout=asyncio.subprocess.DEVNULL,
                stderr=asyncio.subprocess.PIPE,
            )
            _, stderr = await proc.communicate()
            if proc.returncode != 0:
                err = (stderr or b"").decode("utf-8", "replace")[-300:]
                try: await placeholder.edit_text(f"ffmpeg failed: {err}")
                except Exception: pass
                return
        except FileNotFoundError:
            try: await placeholder.edit_text("ffmpeg not available in this build.")
            except Exception: pass
            return
        finally:
            try: os.unlink(vid_path)
            except Exception: pass

        try:
            text, lang = await asyncio.to_thread(_transcribe_sync, wav_path)
        except Exception as e:
            try: await placeholder.edit_text(f"❌ transcription failed: {e}")
            except Exception: pass
            return
        finally:
            try: os.unlink(wav_path)
            except Exception: pass

        caption = (update.message.caption or "").strip()
        if not text and not caption:
            try: await placeholder.edit_text("\U0001F39E пустая транскрипция (немое видео без подписи)")
            except Exception: pass
            return

        body_parts = []
        if text:
            body_parts.append(f"[VIDEO TRANSCRIPT, lang={lang}]\n{text}")
        if caption:
            body_parts.append(f"Caption: {caption}")
        combined = "\n\n".join(body_parts)

        try:
            short_t = text[:200] + ("..." if len(text) > 200 else "") if text else "(no speech)"
            await placeholder.edit_text(f"\U0001F39E [{lang or '?'}] {short_t}")
        except Exception:
            pass

        await _handle_text_inner(update, chat_id, combined)


async def handle_voice(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_allowed(update.effective_user.username): return
    chat_id = update.effective_chat.id
    audio = update.message.voice or update.message.audio
    if not audio: return

    logger.info(
        f"Voice message: chat_id={chat_id} duration={audio.duration}s "
        f"size={audio.file_size}B"
    )

    await _ack_if_locked(update.message, chat_id)
    async with _chat_lock(chat_id):
        try:
            tg_file = await audio.get_file()
            with tempfile.NamedTemporaryFile(suffix=".ogg", delete=False) as tmp:
                tmp_path = tmp.name
            await tg_file.download_to_drive(tmp_path)
        except Exception as e:
            await update.message.reply_text(f"Не смог скачать голосовое: {e}")
            return

        placeholder = await update.message.reply_text("🎙 транскрибирую…")
        try:
            text, lang = await asyncio.to_thread(_transcribe_sync, tmp_path)
        except Exception as e:
            try: await placeholder.edit_text(f"❌ transcription failed: {e}")
            except Exception: pass
            return
        finally:
            try: os.unlink(tmp_path)
            except Exception: pass

        if not text:
            try: await placeholder.edit_text("❌ пустая транскрипция")
            except Exception: pass
            return

        try: await placeholder.edit_text(f"🎙 [{lang}] {text}")
        except Exception: pass

        # Run the transcribed text through the same intent-detection +
        # processing path as a typed message would. We're already in the lock.
        await _handle_text_inner(update, chat_id, text)


_LI_PICK_PATTERN = re.compile(r'^\s*(\d+[ABCDabcd]|EDIT:.+|SKIP)\s*$', re.IGNORECASE)
_LI_INBOX_PATH = r'C:\Users\tonko\OneDrive\Documents\linkedin-engage\inbox.jsonl'


async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_allowed(update.effective_user.username): return
    chat_id = update.effective_chat.id
    user_text = update.message.text
    if not user_text: return

    # LinkedIn-engage pre-filter: route picks to file, skip AI handler
    _msg_text = user_text.strip()
    if _LI_PICK_PATTERN.match(_msg_text):
        try:
            with open(_LI_INBOX_PATH, 'a', encoding='utf-8') as _f:
                _f.write(json.dumps({
                    'ts': datetime.utcnow().isoformat(),
                    'message_id': update.message.message_id,
                    'chat_id': update.message.chat_id,
                    'text': _msg_text
                }, ensure_ascii=False) + '\n')
            await update.message.reply_text(f"ack: {_msg_text} queued for LI engage")
        except Exception as _e:
            await update.message.reply_text(f"li-engage write failed: {_e}")
        return

    fwd_username = None
    try:
        if update.message.forward_origin:
            origin = update.message.forward_origin
            if hasattr(origin, "sender_user") and origin.sender_user:
                u = origin.sender_user
                fwd_name = f"{u.first_name or ''} {u.last_name or ''}".strip()
                fwd_username = u.username
                if u.username: fwd_name += f" (@{u.username})"
                user_text = f"[FORWARDED from {fwd_name}]\n{user_text}"
    except: pass

    await _ack_if_locked(update.message, chat_id)
    async with _chat_lock(chat_id):
        await _handle_text_inner(update, chat_id, user_text, fwd_username)


async def _handle_text_inner(update, chat_id, user_text, fwd_username=None):
    extra = ""
    tl = user_text.lower()
    if GOOGLE_REFRESH_TOKEN:
        cal_kw = ["календар", "расписан", "встреч", "событ", "schedule", "calendar", "meeting", "что у меня", "свободен", "занят", "план на"]
        mail_kw = ["почт", "письм", "email", "mail", "inbox", "gmail", "непрочитан", "написал"]
        drive_kw = ["файл", "документ", "drive", "диск", "найди файл", "doc", "presentation", "таблиц", "sheet"]
        if any(kw in tl for kw in cal_kw):
            ev, _ = get_calendar_events(3, 15)
            if ev: extra += "\n\n" + format_calendar_for_claude(ev)
        if any(kw in tl for kw in mail_kw):
            em, _ = get_emails("is:unread", 10)
            if em: extra += "\n\n" + format_emails_for_claude(em)
        if any(kw in tl for kw in drive_kw):
            df, _ = drive_list_recent(10)
            if df: extra += "\n\n" + format_drive_for_claude(df)
    if HUBSPOT_API_KEY:
        task_kw = ["таск", "задач", "task", "что делать", "что на сегодня", "просроч", "overdue", "to do", "todo"]
        if any(kw in tl for kw in task_kw):
            tasks, _ = get_hubspot_tasks()
            if tasks: extra += "\n\n" + format_tasks_for_claude(tasks)

        # Contact search: detect ANY search intent with a name
        # "найди Олега", "все Викторы", "кто такой Иван", "покажи контакт Петров", etc.
        contact_kw = ["найди", "найти", "поищи", "покажи", "ищи", "search", "find",
                      "контакт", "клиент", "кто такой", "кто такая", "все ", "всех ",
                      "hubspot", "хабспот", "crm", "в базе"]
        if any(kw in tl for kw in contact_kw):
            # Extract the actual search query by removing command words
            search_q = user_text.strip()
            remove_words = ["найди", "найти", "поищи", "покажи", "ищи", "всех", "все", "всю",
                           "контакт", "контакты", "контактов", "клиент", "клиентов", "клиенты",
                           "в хабспот", "в хабспоте", "в crm", "в базе", "в hubspot",
                           "с именем", "по имени", "кто такой", "кто такая",
                           "find", "search", "show", "get", "мне", "мои"]
            q = search_q.lower()
            for w in sorted(remove_words, key=len, reverse=True):
                q = q.replace(w, " ")
            q = " ".join(q.split()).strip()
            # Clean up trailing/leading punctuation
            q = q.strip("?!.,;: ")
            if q and len(q) > 1 and not q.isdigit():
                contacts, _ = search_hubspot_contacts(q)
                if contacts is not None:
                    extra += "\n\n" + format_contacts_for_claude(contacts)

        # Deal search
        deal_kw = ["сделк", "deal", "pipeline", "воронк"]
        if any(kw in tl for kw in deal_kw):
            deals, _ = search_hubspot_deals()
            if deals: extra += "\n\n" + format_deals_for_claude(deals)

    url_match = re.search(r'https?://\S+', user_text)
    if url_match:
        fetched = fetch_url_text(url_match.group(0))
        extra += f"\n\n[URL CONTENT]\n{fetched}"
    elif any(kw in tl for kw in ["linkedin.com/in/", "linkedin профил", "linkedin profile"]):
        ln_match = re.search(r'linkedin\.com/in/[\w-]+', user_text)
        if ln_match:
            fetched = fetch_url_text(f"https://www.{ln_match.group(0)}")
            extra += f"\n\n[LINKEDIN PROFILE]\n{fetched}"
    elif any(kw in tl for kw in ["реддит", "reddit", "r/popular", "r/all"]):
        reddit_match = re.search(r'r/(\w+)', user_text)
        if reddit_match:
            fetched = fetch_url_text(f"https://www.reddit.com/r/{reddit_match.group(1)}/")
        else:
            fetched = fetch_url_text("https://www.reddit.com/r/popular/")
        extra += f"\n\n[REDDIT RSS]\n{fetched}"

    schedule_kw = ["каждый день в", "каждое утро", "каждый вечер", "ежедневно в", "every day at", "every morning"]
    cancel_kw = ["отмени задачу", "удали задачу", "список задач", "cancel job", "my jobs"]
    if any(kw in tl for kw in schedule_kw):
        reply = await handle_schedule_command(update, chat_id, user_text)
        await update.message.reply_text(md_to_html(reply), parse_mode="HTML")
        return
    if any(kw in tl for kw in cancel_kw):
        reply = await handle_cancel_schedule(update, chat_id, user_text)
        await update.message.reply_text(md_to_html(reply), parse_mode="HTML")
        return

    # Garage trigger: intercept before HubSpot/calendar enrichment matters
    _is_garage_r, _garage_payload = _detect_garage_railway(user_text)
    if _is_garage_r:
        garage_text = _GARAGE_SYSTEM_CONTEXT_RAILWAY + (_garage_payload if _garage_payload else "Garage agent here. What do you need?")
        logger.info("chat %s: garage trigger detected on Railway bot", chat_id)
        await _process_message(update, chat_id, garage_text, fwd_username=fwd_username)
        return

    await _process_message(update, chat_id, user_text + extra, fwd_username=fwd_username)


async def _keep_typing(bot, chat_id, stop_event):
    while not stop_event.is_set():
        await bot.send_chat_action(chat_id=chat_id, action="typing")
        await asyncio.sleep(4)


async def _process_message(update, chat_id, content, fwd_username=None):
    _convdb.db_append(chat_id, "user", content)
    _convdb.db_truncate(chat_id, MAX_HISTORY)
    history = _convdb.db_load(chat_id)
    stop_typing = asyncio.Event()
    typing_task = asyncio.create_task(
        _keep_typing(update.get_bot(), chat_id, stop_typing)
    )
    try:
        response = await asyncio.wait_for(
            asyncio.get_event_loop().run_in_executor(
                None,
                lambda: client.messages.create(
                    model=MODEL,
                    max_tokens=4096,
                    system=SYSTEM_PROMPT,
                    messages=history,
                    tools=[{"type": "web_search_20250305", "name": "web_search"}]
                )
            ),
            timeout=60
        )
        stop_typing.set()
        await typing_task
        reply = "\n".join(
            block.text for block in response.content
            if hasattr(block, "text")
        )
        _convdb.db_append(chat_id, "assistant", reply)

        # Check if Claude wants to search HubSpot first (two-phase: search → act)
        hs_search_q = extract_tag_text(reply, "hubspot_search")
        if hs_search_q:
            contacts, _ = search_hubspot_contacts(hs_search_q.strip())
            search_data = format_contacts_for_claude(contacts) if contacts else "[HUBSPOT CONTACTS]\nNo contacts found for '" + hs_search_q + "'."
            # Don't show intermediate reply to user — replace last assistant turn
            _convdb.db_replace_last(chat_id, "assistant", clean_response(reply) or "Searching...")
            follow_up = f"[SEARCH RESULTS for '{hs_search_q}']\n{search_data}\n\nNow complete the user's original request using the contact_id from results above. Do NOT output another hubspot_search tag."
            if contacts:
                c = contacts[0]
                follow_up += f"\nFirst match: ID={c['id']}, {c['properties'].get('firstname','')} {c['properties'].get('lastname','')}"
            _convdb.db_append(chat_id, "user", follow_up)
            history = _convdb.db_load(chat_id)
            response2 = await asyncio.wait_for(
                asyncio.get_event_loop().run_in_executor(
                    None,
                    lambda: client.messages.create(
                        model=MODEL, max_tokens=4096, system=SYSTEM_PROMPT,
                        messages=history,
                    )
                ), timeout=60
            )
            reply = "\n".join(block.text for block in response2.content if hasattr(block, "text"))
            _convdb.db_append(chat_id, "assistant", reply)

        hs_create = extract_tag_json(reply, "hubspot_create")
        hs_update = extract_tag_json(reply, "hubspot_update")
        tg_username = extract_tag_text(reply, "hubspot_contact") or fwd_username
        if tg_username: tg_username = tg_username.lstrip("@")
        hs_task = extract_tag_json(reply, "hubspot_task")
        hs_edit = extract_tag_json(reply, "hubspot_edit")
        hs_deal = extract_tag_json(reply, "hubspot_deal")
        hs_complete = extract_tag_text(reply, "hubspot_complete_task")
        cal_create = extract_tag_json(reply, "calendar_create")
        email_send = extract_tag_json(reply, "gmail_send")
        email_draft = extract_tag_json(reply, "gmail_draft")
        # Capture [ATTACH:] paths before clean_response strips them from text
        _attach_paths = [m.group(1).strip() for m in ATTACH_RE.finditer(reply or "")]
        clean = clean_response(reply)

        if hs_create:
            # Extract telegram username from website field for cache key
            website = hs_create.get("website", "")
            tg_match = re.search(r't\.me/(\w+)', website)
            create_username = tg_match.group(1).lower() if tg_match else None

            # Duplicate check: if contact already in cache, offer update instead of create
            name = f"{hs_create.get('firstname', '')} {hs_create.get('lastname', '')}".strip() or "—"
            dup_cid = created_contacts.get(create_username) if create_username else None
            if not dup_cid and name != "—":
                dup_cid = created_contacts.get(name.lower())
            if dup_cid:
                link = f"https://app.hubspot.com/contacts/47345195/record/0-1/{dup_cid}"
                await update.message.reply_text(
                    f"{md_to_html(clean)}\n\n⚠️ Контакт <b>{esc(name)}</b> уже создан в этой сессии.\n<a href=\"{link}\">Открыть в HubSpot</a>",
                    parse_mode="HTML", disable_web_page_preview=True)
            else:
                pending_updates[chat_id] = {
                    "type": "hubspot_create",
                    "data": hs_create,
                    "tg_username": create_username,
                }
                email_val = hs_create.get("email", "—")
                phone_val = hs_create.get("phone", "")
                job_val = hs_create.get("jobtitle", "") or hs_create.get("job_title", "")
                company_val = hs_create.get("company", "")
                stage = LIFECYCLE_STAGES.get(hs_create.get("lifecyclestage", ""), hs_create.get("lifecyclestage", "—"))
                details = f"Имя: {esc(name)}\nEmail: {esc(email_val)}"
                if phone_val: details += f"\nТел: {esc(phone_val)}"
                if job_val: details += f"\nДолжность: {esc(job_val)}"
                if company_val: details += f"\nКомпания: {esc(company_val)}"
                details += f"\nStage: {esc(stage)}"
                if website: details += f"\nTelegram: {esc(website)}"
                kb = [[InlineKeyboardButton("✅ Создать", callback_data="hsc_confirm"), InlineKeyboardButton("❌ Отмена", callback_data="hsc_cancel")]]
                await update.message.reply_text(
                    f"{md_to_html(clean)}\n\n━━━━━━━━━━━━━━━\n<b>Создать контакт в HubSpot:</b>\n{details}",
                    parse_mode="HTML", reply_markup=InlineKeyboardMarkup(kb))

        elif hs_update and tg_username:
            # If this username is in the local cache (just created), apply update directly without search
            if tg_username.lower() in created_contacts:
                cid = created_contacts[tg_username.lower()]
                note = hs_update.get("suggested_note")
                props = {}
                lc = hs_update.get("suggested_lifecycle")
                ls = hs_update.get("suggested_lead_status")
                if lc: props["lifecyclestage"] = lc
                if ls: props["hs_lead_status"] = ls
                if props:
                    update_contact(cid, props)
                if note:
                    add_note_to_contact(cid, note)
                link = f"https://app.hubspot.com/contacts/47345195/record/0-1/{cid}"
                await update.message.reply_text(
                    f"{md_to_html(clean)}\n\n✅ Данные добавлены в контакт.\n<a href=\"{link}\">HubSpot</a>",
                    parse_mode="HTML", disable_web_page_preview=True)
            else:
                await _send_hubspot_update(update, chat_id, hs_update, tg_username, clean)

        elif cal_create:
            pending_updates[chat_id] = {"type": "calendar", "data": cal_create}
            kb = [[InlineKeyboardButton("✅ Создать", callback_data="cal_confirm"), InlineKeyboardButton("❌ Отмена", callback_data="cal_cancel")]]
            await update.message.reply_text(
                f"{md_to_html(clean)}\n\n——————————\n<b>Создать событие:</b>\n"
                f"{esc(cal_create.get('summary',''))}\n{esc(cal_create.get('start',''))} — {esc(cal_create.get('end',''))}",
                parse_mode="HTML", reply_markup=InlineKeyboardMarkup(kb))
        elif email_send:
            pending_updates[chat_id] = {"type": "email", "data": email_send}
            kb = [[InlineKeyboardButton("✅ Отправить", callback_data="email_confirm"), InlineKeyboardButton("❌ Отмена", callback_data="email_cancel")]]
            await update.message.reply_text(
                f"{md_to_html(clean)}\n\n——————————\n<b>Письмо:</b>\nTo: {esc(email_send.get('to',''))}\n"
                f"Subject: {esc(email_send.get('subject',''))}\n{esc(email_send.get('body','')[:200])}...",
                parse_mode="HTML", reply_markup=InlineKeyboardMarkup(kb))
        elif email_draft:
            pending_updates[chat_id] = {"type": "draft", "data": email_draft}
            kb = [[InlineKeyboardButton("💾 Сохранить черновик", callback_data="draft_confirm"), InlineKeyboardButton("❌ Отмена", callback_data="draft_cancel")]]
            await update.message.reply_text(
                f"{md_to_html(clean)}\n\n——————————\n<b>Черновик:</b>\nTo: {esc(email_draft.get('to','(не указан)'))}\n"
                f"Subject: {esc(email_draft.get('subject',''))}\n{esc(email_draft.get('body','')[:200])}...",
                parse_mode="HTML", reply_markup=InlineKeyboardMarkup(kb))
        elif hs_task:
            # Create HubSpot task immediately (no confirmation needed for tasks)
            result = create_hubspot_task(
                subject=hs_task.get("subject", "Task"),
                body=hs_task.get("body", ""),
                due_date=hs_task.get("due_date"),
                priority=hs_task.get("priority", "MEDIUM"),
                contact_id=hs_task.get("contact_id"),
            )
            if "error" not in result:
                tid = result.get("id", "?")
                await update.message.reply_text(
                    f"{md_to_html(clean)}\n\n✅ Задача создана в HubSpot (ID: {tid})",
                    parse_mode="HTML")
            else:
                await update.message.reply_text(
                    f"{md_to_html(clean)}\n\n❌ Ошибка: {esc(str(result.get('message', result.get('error', '?')))[:200])}",
                    parse_mode="HTML")

        elif hs_edit:
            # Edit contact fields: {"search": "Виктор", "updates": {"email": "x@y.com"}}
            search_q = hs_edit.get("search", "")
            updates = hs_edit.get("updates", {})
            if search_q and updates:
                result = edit_hubspot_contact(search_q, updates)
                if "error" not in result:
                    name = result.get("_found_name", search_q)
                    cid = result.get("_contact_id", "")
                    link = f"https://app.hubspot.com/contacts/47345195/record/0-1/{cid}"
                    fields = ", ".join(f"{k}={v}" for k, v in updates.items())
                    await update.message.reply_text(
                        f"{md_to_html(clean)}\n\n✅ Обновлён <b>{esc(name)}</b>: {esc(fields)}\n<a href=\"{link}\">HubSpot</a>",
                        parse_mode="HTML", disable_web_page_preview=True)
                else:
                    await update.message.reply_text(
                        f"{md_to_html(clean)}\n\n❌ {esc(str(result.get('error', '?'))[:200])}",
                        parse_mode="HTML")
            else:
                await _send_reply(update, clean)

        elif hs_deal:
            result = create_hubspot_deal(
                name=hs_deal.get("name", "Deal"),
                stage=hs_deal.get("stage", "appointmentscheduled"),
                amount=hs_deal.get("amount"),
                contact_id=hs_deal.get("contact_id"),
            )
            if "error" not in result:
                did = result.get("id", "?")
                await update.message.reply_text(
                    f"{md_to_html(clean)}\n\n✅ Сделка создана (ID: {did})",
                    parse_mode="HTML")
            else:
                await update.message.reply_text(
                    f"{md_to_html(clean)}\n\n❌ Ошибка: {esc(str(result.get('message', '?'))[:200])}",
                    parse_mode="HTML")

        elif hs_complete:
            result = complete_hubspot_task(hs_complete.strip())
            if "error" not in result:
                await update.message.reply_text(f"{md_to_html(clean)}\n\n✅ Задача закрыта.", parse_mode="HTML")
            else:
                await update.message.reply_text(f"{md_to_html(clean)}\n\n❌ Ошибка: {esc(str(result)[:200])}", parse_mode="HTML")

        else:
            await _send_reply(update, clean)

        # After whatever branch ran, push any [ATTACH:] files Claude requested
        if _attach_paths:
            await _send_attachments(update.message, _attach_paths)
    except asyncio.TimeoutError:
        stop_typing.set()
        await typing_task
        logger.error("Request timed out after 60s")
        await update.message.reply_text("⏱ Запрос занял слишком долго (>60 сек). Попробуй ещё раз.")
    except Exception as e:
        stop_typing.set()
        await typing_task
        logger.error(f"Error: {e}")
        await update.message.reply_text(f"❌ Ошибка: {esc(str(e)[:200])}")


_MD_FENCE_RE = re.compile(r"```(\w*)\n?(.*?)```", re.DOTALL)
_MD_INLINE_CODE_RE = re.compile(r"`([^`\n]+)`")
_MD_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)\s]+)\)")
_MD_BOLD_RE = re.compile(r"\*\*(.+?)\*\*", re.DOTALL)
_MD_ITALIC_STAR_RE = re.compile(r"(?<![\*\w])\*(?!\s)([^*\n]+?)(?<!\s)\*(?![\*\w])")
_MD_ITALIC_UND_RE = re.compile(r"(?<!\w)_(?!\s)([^_\n]+?)(?<!\s)_(?!\w)")
_MD_HEADER_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$", re.MULTILINE)
_MD_HR_RE = re.compile(r"^[ \t]*(-{3,}|\*{3,})[ \t]*$", re.MULTILINE)
_MD_BULLET_RE = re.compile(r"^(\s*)[-*]\s+", re.MULTILINE)


def md_to_html(text):
    """Convert Claude's GitHub-flavored Markdown to Telegram-safe HTML.

    Handles: fenced code blocks, inline code, links, bold, italic, headers,
    horizontal rules, bullet lists. HTML-escapes <, >, & in non-tag text.
    """
    if not text:
        return ""

    fences = []

    def _fence_sub(m):
        idx = len(fences)
        fences.append((m.group(1) or "", m.group(2) or ""))
        return f"\x00F{idx}\x00"
    text = _MD_FENCE_RE.sub(_fence_sub, text)

    inlines = []

    def _inline_sub(m):
        idx = len(inlines)
        inlines.append(m.group(1))
        return f"\x00I{idx}\x00"
    text = _MD_INLINE_CODE_RE.sub(_inline_sub, text)

    links = []

    def _link_sub(m):
        idx = len(links)
        links.append((m.group(1), m.group(2)))
        return f"\x00L{idx}\x00"
    text = _MD_LINK_RE.sub(_link_sub, text)

    # Escape HTML specials in remaining text
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    # Headers → bold on own line
    text = _MD_HEADER_RE.sub(lambda m: f"<b>{m.group(2).strip()}</b>", text)

    # Horizontal rule
    text = _MD_HR_RE.sub("─────────────", text)

    # Bold before italic
    text = _MD_BOLD_RE.sub(r"<b>\1</b>", text)
    text = _MD_ITALIC_STAR_RE.sub(r"<i>\1</i>", text)
    text = _MD_ITALIC_UND_RE.sub(r"<i>\1</i>", text)

    # Bullets
    text = _MD_BULLET_RE.sub(r"\1• ", text)

    # Restore inline code
    for i, code in enumerate(inlines):
        safe = code.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        text = text.replace(f"\x00I{i}\x00", f"<code>{safe}</code>")

    # Restore links
    for i, (label, url) in enumerate(links):
        safe_label = label.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        safe_url = url.replace("&", "&amp;").replace('"', "&quot;")
        text = text.replace(f"\x00L{i}\x00", f'<a href="{safe_url}">{safe_label}</a>')

    # Restore fenced code blocks
    for i, (lang, code) in enumerate(fences):
        safe_code = code.rstrip("\n").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if lang:
            repl = f'<pre><code class="language-{lang}">{safe_code}</code></pre>'
        else:
            repl = f"<pre>{safe_code}</pre>"
        text = text.replace(f"\x00F{i}\x00", repl)

    return text.strip()


async def _send_reply(update, text):
    if not text: text = "(пустой ответ)"
    text = md_to_html(text)
    # Remove any leftover XML-like tags that Telegram would choke on
    text = re.sub(r'<(?!/?(?:b|i|u|s|code|pre|a)\b)[^>]+>', '', text)
    for i in range(0, len(text), 4096):
        try:
            await update.message.reply_text(text[i:i+4096], parse_mode="HTML")
        except Exception:
            # Fallback: send without HTML parsing
            await update.message.reply_text(text[i:i+4096])


async def _send_hubspot_update(update, chat_id, hs_update, tg_username, clean):
    contacts = search_contact_by_telegram(tg_username)
    if contacts:
        c = contacts[0]
        p = c.get("properties", {})
        name = f"{p.get('firstname', '')} {p.get('lastname', '')}".strip() or tg_username
        cid = c["id"]
        link = f"https://app.hubspot.com/contacts/47345195/record/0-1/{cid}"
        cs = LIFECYCLE_STAGES.get(p.get("lifecyclestage", ""), "—")
        cst = LEAD_STATUSES.get(p.get("hs_lead_status", ""), "—")
        ns = LIFECYCLE_STAGES.get(hs_update.get("suggested_lifecycle", ""), "—")
        nst = LEAD_STATUSES.get(hs_update.get("suggested_lead_status", ""), "—")
        pending_updates[chat_id] = {"type": "hubspot", "contact_id": cid, "contact_name": name,
            "lifecycle": hs_update.get("suggested_lifecycle"), "lead_status": hs_update.get("suggested_lead_status"), "note": hs_update.get("suggested_note")}
        kb = [[InlineKeyboardButton("✅ Всё", callback_data="hs_confirm"), InlineKeyboardButton("❌", callback_data="hs_cancel")],
              [InlineKeyboardButton("📝 Заметку", callback_data="hs_note_only"), InlineKeyboardButton("📊 Статус", callback_data="hs_status_only")]]
        await update.message.reply_text(
            f"{md_to_html(clean)}\n\n━━━━━━━━━━━━━━━\n<b>HubSpot: {esc(name)}</b>\n\n"
            f"Stage: {esc(cs)} → <b>{esc(ns)}</b>\nStatus: {esc(cst)} → <b>{esc(nst)}</b>\n"
            f"Заметка: {esc(hs_update.get('suggested_note', '—'))}\n\n<b>Нажми кнопку:</b>",
            parse_mode="HTML", reply_markup=InlineKeyboardMarkup(kb))
    else:
        # No search spam — just inform and suggest creating
        await update.message.reply_text(
            f"{md_to_html(clean)}\n\nКонтакт <b>{esc(tg_username)}</b> не найден в HubSpot.\n"
            f"Напиши <b>создай контакт</b> чтобы добавить.",
            parse_mode="HTML")


async def handle_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    q = update.callback_query
    await q.answer()
    chat_id = q.message.chat_id
    pending = pending_updates.get(chat_id)
    if not pending:
        await q.edit_message_text("Нет активного действия.")
        return

    t = pending.get("type", "hubspot")

    if t == "hubspot_create":
        if q.data == "hsc_confirm":
            r = create_hubspot_contact(dict(pending["data"]))
            if "error" not in r:
                cid = r["id"]
                link = f"https://app.hubspot.com/contacts/47345195/record/0-1/{cid}"
                d = pending["data"]
                name = f"{d.get('firstname', '')} {d.get('lastname', '')}".strip() or "—"
                # Cache by tg username AND by full name to bypass HubSpot search index delay
                tg_u = pending.get("tg_username")
                if tg_u:
                    created_contacts[tg_u.lower()] = cid
                # Also cache by name and email so updates find it despite HubSpot index delay
                if name and name != "—":
                    created_contacts[name.lower()] = cid
                email = d.get("email", "")
                if email:
                    created_contacts[email.lower()] = cid
                await q.edit_message_text(
                    f"✅ Контакт <b>{esc(name)}</b> создан!\n<a href=\"{link}\">Открыть в HubSpot</a>",
                    parse_mode="HTML", disable_web_page_preview=True)
            else:
                err_msg = str(r.get("message", "") or r.get("error", ""))
                m = re.search(r"Existing ID:\s*(\d+)", err_msg)
                if m:
                    existing_cid = m.group(1)
                    existing_link = f"https://app.hubspot.com/contacts/47345195/record/0-1/{existing_cid}"
                    d = pending["data"]
                    name = f"{d.get('firstname', '')} {d.get('lastname', '')}".strip() or "—"
                    tg_u = pending.get("tg_username")
                    if tg_u:
                        created_contacts[tg_u.lower()] = existing_cid
                    if name and name != "—":
                        created_contacts[name.lower()] = existing_cid
                    email = d.get("email", "")
                    if email:
                        created_contacts[email.lower()] = existing_cid
                    note = d.get("notes_initial")
                    note_suffix = ""
                    if note:
                        nr = add_note_to_contact(existing_cid, note)
                        note_suffix = "\n📝 Заметка добавлена." if "error" not in nr else ""
                    await q.edit_message_text(
                        f"⚠️ <b>{esc(name)}</b> уже есть в HubSpot.\n<a href=\"{existing_link}\">Открыть карточку</a>{note_suffix}",
                        parse_mode="HTML", disable_web_page_preview=True)
                else:
                    await q.edit_message_text(f"Ошибка создания: {err_msg[:300]}")
        else:
            await q.edit_message_text("Отменено.")

    elif t == "hubspot":
        cid = pending["contact_id"]
        name = pending["contact_name"]
        link = f"https://app.hubspot.com/contacts/47345195/record/0-1/{cid}"
        if q.data == "hs_confirm":
            props = {}
            if pending.get("lifecycle"): props["lifecyclestage"] = pending["lifecycle"]
            if pending.get("lead_status"): props["hs_lead_status"] = pending["lead_status"]
            r = update_contact(cid, props) if props else {}
            nr = add_note_to_contact(cid, pending["note"]) if pending.get("note") else {}
            if "error" not in r and "error" not in nr:
                await q.edit_message_text(f"✅ <b>{esc(name)}</b> обновлён!\n<a href=\"{link}\">HubSpot</a>", parse_mode="HTML", disable_web_page_preview=True)
            else:
                await q.edit_message_text(f"Ошибка: {(r.get('message','') or nr.get('message',''))[:500]}")
        elif q.data == "hs_note_only":
            r = add_note_to_contact(cid, pending["note"]) if pending.get("note") else {"error": "no note"}
            if "error" not in r:
                await q.edit_message_text(f"📝 Заметка для <b>{esc(name)}</b>\n<a href=\"{link}\">HubSpot</a>", parse_mode="HTML", disable_web_page_preview=True)
            else: await q.edit_message_text(f"Ошибка: {r.get('message','')[:500]}")
        elif q.data == "hs_status_only":
            props = {}
            if pending.get("lifecycle"): props["lifecyclestage"] = pending["lifecycle"]
            if pending.get("lead_status"): props["hs_lead_status"] = pending["lead_status"]
            if props:
                r = update_contact(cid, props)
                if "error" not in r:
                    await q.edit_message_text(f"📊 <b>{esc(name)}</b> обновлён\n<a href=\"{link}\">HubSpot</a>", parse_mode="HTML", disable_web_page_preview=True)
                else: await q.edit_message_text(f"Ошибка: {r.get('message','')[:500]}")
        elif q.data == "hs_cancel":
            await q.edit_message_text("Отменено.")

    elif t == "calendar":
        if q.data == "cal_confirm":
            d = pending["data"]
            r = create_calendar_event(d["summary"], d["start"], d["end"], d.get("description", ""))
            await q.edit_message_text(f"✅ Событие создано: {d['summary']}" if "error" not in r else f"Ошибка: {(r.get('message') or str(r.get('error') or '?'))[:300]}")
        else: await q.edit_message_text("Отменено.")

    elif t == "email":
        if q.data == "email_confirm":
            d = pending["data"]
            r = send_email(d["to"], d["subject"], d["body"])
            await q.edit_message_text(f"✅ Письмо отправлено: {d['to']}" if "error" not in r else f"Ошибка: {(r.get('message') or str(r.get('error') or '?'))[:300]}")
        else: await q.edit_message_text("Отменено.")

    elif t == "draft":
        if q.data == "draft_confirm":
            d = pending["data"]
            r = save_draft(d.get("to", ""), d["subject"], d["body"])
            await q.edit_message_text("💾 Черновик сохранён в Gmail." if "error" not in r else f"Ошибка: {(r.get('message') or str(r.get('error') or '?'))[:300]}")
        else: await q.edit_message_text("Отменено.")

    pending_updates.pop(chat_id, None)


async def handle_schedule_command(update, chat_id, user_text):
    hour, minute = 9, 0
    match = re.search(r'в\s+(\d{1,2})(?::(\d{2}))?\s*(утра|вечера|am|pm)?', user_text.lower())
    if not match:
        match = re.search(r'at\s+(\d{1,2})(?::(\d{2}))?\s*(am|pm)?', user_text.lower())
    if match:
        hour = int(match.group(1))
        minute = int(match.group(2)) if match.group(2) else 0
        suffix = match.group(3) or ""
        if suffix in ("вечера", "pm") and hour < 12:
            hour += 12

    job_id = f"job_{chat_id}_{hour}_{minute}"
    task = user_text

    if scheduler.get_job(job_id):
        scheduler.remove_job(job_id)
        scheduled_jobs[chat_id] = [j for j in scheduled_jobs[chat_id] if j["id"] != job_id]

    async def scheduled_task():
        try:
            response = client.messages.create(
                model=MODEL,
                max_tokens=2048,
                system=SYSTEM_PROMPT,
                messages=[{"role": "user", "content": task}],
                tools=[{"type": "web_search_20250305", "name": "web_search"}]
            )
            reply = "\n".join(block.text for block in response.content if hasattr(block, "text"))
            reply = md_to_html(reply)
            from telegram import Bot
            bot = Bot(token=TELEGRAM_TOKEN)
            for i in range(0, len(reply), 4096):
                await bot.send_message(chat_id=chat_id, text=reply[i:i+4096], parse_mode="HTML")
        except Exception as e:
            logger.error(f"Scheduled task error: {e}")

    scheduler.add_job(
        scheduled_task,
        CronTrigger(hour=hour, minute=minute, timezone="America/Los_Angeles"),
        id=job_id,
        replace_existing=True
    )
    scheduled_jobs[chat_id].append({"id": job_id, "time": f"{hour:02d}:{minute:02d}", "task": task[:80]})
    return f"✅ Задача создана: каждый день в <b>{hour:02d}:{minute:02d}</b> PT\nЗадание: {esc(task[:100])}\n\nЧтобы отменить: напиши <b>отмени задачу {hour:02d}:{minute:02d}</b>"


async def handle_cancel_schedule(update, chat_id, user_text):
    match = re.search(r'(\d{1,2}):?(\d{2})?', user_text)
    if match:
        hour = int(match.group(1))
        minute = int(match.group(2)) if match.group(2) else 0
        job_id = f"job_{chat_id}_{hour}_{minute}"
        if scheduler.get_job(job_id):
            scheduler.remove_job(job_id)
            scheduled_jobs[chat_id] = [j for j in scheduled_jobs[chat_id] if j["id"] != job_id]
            return f"✅ Задача в {hour:02d}:{minute:02d} отменена."
    jobs = scheduled_jobs.get(chat_id, [])
    if not jobs:
        return "Нет активных задач."
    lines = ["Активные задачи:"]
    for j in jobs:
        lines.append(f"• {j['time']} — {j['task']}")
    return "\n".join(lines)


def main():
    app = (
        Application.builder()
        .token(TELEGRAM_TOKEN)
        .concurrent_updates(True)
        .build()
    )
    for cmd, fn in [("start", start), ("reset", reset), ("debug", debug_cmd), ("cal", calendar_cmd),
                    ("mail", mail_cmd), ("drive", drive_cmd), ("tasks", tasks_cmd), ("model", model_info), ("setmodel", set_model), ("find", find_contact)]:
        app.add_handler(CommandHandler(cmd, fn))
    for i in range(1, 15):
        app.add_handler(CommandHandler(f"cal{i}", calendar_cmd))
    app.add_handler(CallbackQueryHandler(handle_callback))
    app.add_handler(MessageHandler(filters.PHOTO, handle_photo))
    app.add_handler(MessageHandler(filters.VOICE | filters.AUDIO, handle_voice))
    app.add_handler(MessageHandler(filters.VIDEO | filters.VIDEO_NOTE, handle_video))
    app.add_handler(MessageHandler(filters.Document.ALL, handle_document))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))
    logger.info(f"Bot started. Model: {MODEL} | Google: {'YES' if GOOGLE_REFRESH_TOKEN else 'NO'} | HubSpot: {'YES' if HUBSPOT_API_KEY else 'NO'}")
    async def post_init(application):
        scheduler.start()

    app.post_init = post_init
    app.run_polling()

if __name__ == "__main__":
    main()
